from base64 import urlsafe_b64encode
from pathlib import Path
from textwrap import wrap

import qrcode
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


GENERAL = Path(__file__).resolve().parents[1]
DOCUMENTS = GENERAL / "Documentos_Academicos"
ASSETS = GENERAL / "Recursos_Visuales"
OUT = DOCUMENTS / "Ensayo_1_Origen_y_Evolucion_Ciencia_de_Datos_Marlenis_FINAL_APA_UASD.docx"
LOGO = ASSETS / "uasd_logo.png"
TIMELINE = ASSETS / "figura_1_linea_tiempo_dsc.png"
QR = ASSETS / "anexo_qr_linea_tiempo_dsc.png"

BLUE = "0073B7"
TEAL = "1B9AAA"
GOLD = "F2B705"
LIGHT_BLUE = "EAF6FC"
LIGHT_GOLD = "FFF4CC"
GRAY = "F7F7F7"
BLACK = "111111"


def set_margins(section):
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def set_page_numbering(section, fmt="decimal", start=1):
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))
    pg_num_type.set(qn("w:fmt"), fmt)


def add_page_field(paragraph, roman=False):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE \\* roman" if roman else "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def configure_header(section, roman=False):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header_p = section.header.paragraphs[0]
    header_p.text = ""
    add_page_field(header_p, roman=roman)
    footer_p = section.footer.paragraphs[0]
    footer_p.text = ""


def clear_header_footer(section):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    section.header.paragraphs[0].text = ""
    section.footer.paragraphs[0].text = ""


def keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_text(cell, text, bold=False, size=11, color=BLACK):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    r.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def style_doc(doc):
    set_margins(doc.sections[0])
    clear_header_footer(doc.sections[0])
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        style.font.bold = True
        style.paragraph_format.line_spacing = 2
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def para(doc, text="", align=None, bold=False, italic=False, indent=True, size=12, color=BLACK):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.5)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    r.bold = bold
    r.italic = italic
    return p


def h1(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(BLACK)
    r.bold = True
    keep_with_next(p)
    return p


def h2(doc, text):
    p = doc.add_paragraph(style="Heading 2")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(BLACK)
    r.bold = True
    keep_with_next(p)
    return p


def h3(doc, text):
    p = doc.add_paragraph(style="Heading 3")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(BLACK)
    r.bold = True
    r.italic = True
    keep_with_next(p)
    return p


def add_updateable_toc(doc):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "Haga clic derecho y seleccione Actualizar campo para actualizar la tabla de contenido."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(placeholder)
    run._r.append(end)


def page_break(doc):
    doc.add_page_break()


def generate_timeline_image():
    width, height = 1600, 720
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype("Arial.ttf", 44)
        year_font = ImageFont.truetype("Arial Bold.ttf", 26)
        text_font = ImageFont.truetype("Arial.ttf", 24)
    except OSError:
        title_font = year_font = text_font = ImageFont.load_default()

    draw.rectangle([0, 0, width, 95], fill=(0, 115, 183))
    draw.text((60, 28), "Evolucion de la ciencia de datos", fill="white", font=title_font)
    y = 360
    draw.line([95, y, width - 95, y], fill=(0, 115, 183), width=8)
    events = [
        ("1962", "Tukey propone\nel futuro del\nanalisis de datos"),
        ("1989", "Primer taller\nKDD"),
        ("1996", "Fayyad et al.\nformalizan KDD"),
        ("1997", "Revista Data\nMining and\nKnowledge Discovery"),
        ("2001", "Cleveland usa\nData Science"),
        ("2001", "Breiman plantea\nlas dos culturas"),
        ("2017", "Donoho reconstruye\n50 anos de\nciencia de datos"),
    ]
    xs = [115, 340, 565, 790, 1015, 1240, 1465]
    for i, ((year, label), x) in enumerate(zip(events, xs)):
        color = (242, 183, 5) if i % 2 == 0 else (27, 154, 170)
        draw.ellipse([x - 28, y - 28, x + 28, y + 28], fill=color, outline=(17, 17, 17), width=3)
        box_top = 145 if i % 2 == 0 else 430
        box_bottom = box_top + 165
        draw.rounded_rectangle([x - 98, box_top, x + 98, box_bottom], radius=18, fill=(245, 248, 250), outline=color, width=4)
        draw.text((x - 42, box_top + 16), year, fill=(0, 65, 105), font=year_font)
        lines = label.split("\n")
        ty = box_top + 55
        for line in lines:
            draw.text((x - 82, ty), line, fill=(17, 17, 17), font=text_font)
            ty += 28
        if i % 2 == 0:
            draw.line([x, box_bottom, x, y - 31], fill=color, width=4)
        else:
            draw.line([x, y + 31, x, box_top], fill=color, width=4)
    draw.text((60, 665), "Nota. Elaboracion propia con base en Tukey (1962), Fayyad et al. (1996), Cleveland (2001), Breiman (2001) y Donoho (2017).", fill=(80, 80, 80), font=text_font)
    img.save(TIMELINE)


def generate_qr():
    mermaid = """timeline
    title Evolucion de la Ciencia de Datos
    1962 : Tukey propone el futuro del analisis de datos
    1989 : Primer taller KDD
    1996 : Fayyad et al. formalizan el proceso KDD
    1997 : Data Mining and Knowledge Discovery
    2001 : Cleveland propone Data Science
    2001 : Breiman contrasta dos culturas del modelado
    2017 : Donoho reconstruye 50 anos de ciencia de datos"""
    url = f"https://mermaid.ink/img/{urlsafe_b64encode(mermaid.encode('utf-8')).decode('ascii')}"
    qrcode.make(url).save(QR)
    return url


def title_page(doc):
    for _ in range(2):
        para(doc, "", indent=False)
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(1.35))
    para(doc, "Presentación", WD_ALIGN_PARAGRAPH.CENTER, bold=True, indent=False, color=BLUE, size=14)
    para(doc, "Universidad Autónoma de Santo Domingo (UASD)", WD_ALIGN_PARAGRAPH.CENTER, bold=True, indent=False, color=BLUE, size=15)
    para(doc, "Facultad de Ciencias", WD_ALIGN_PARAGRAPH.CENTER, bold=True, indent=False, color=TEAL)
    para(doc, "Programa de Maestría", WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    para(doc, "Ensayo 1", WD_ALIGN_PARAGRAPH.CENTER, bold=True, indent=False, color=BLUE, size=14)
    para(doc, "Origen y Evolución de la Ciencia de Datos", WD_ALIGN_PARAGRAPH.CENTER, bold=True, indent=False, size=14)
    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    rows = [
        ("Asignatura", "INF-8237 Ciencia de Datos I"),
        ("Participante", "Marlenis Judith Concepción Cuevas"),
        ("Maestro", "Silverio del Orbe"),
        ("Periodo", "17 de mayo al 28 de junio de 2026"),
        ("Horario", "Domingos, 9:00 a.m. a 5:00 p.m."),
        ("Fecha de entrega", "23 de mayo de 2026"),
    ]
    for i, (left, right) in enumerate(rows):
        set_cell_shading(table.cell(i, 0), LIGHT_BLUE if i % 2 == 0 else LIGHT_GOLD)
        set_cell_shading(table.cell(i, 1), GRAY if i % 2 == 0 else "FFFFFF")
        cell_text(table.cell(i, 0), left, bold=True, color=BLUE)
        cell_text(table.cell(i, 1), right)
    para(doc, "Santo Domingo, República Dominicana", WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    prelim_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(prelim_section)
    set_page_numbering(prelim_section, "lowerRoman", 1)
    configure_header(prelim_section, roman=True)


def abstract_pages(doc):
    h1(doc, "Resumen")
    para(doc, "La ciencia de datos es un campo interdisciplinario que surge de la convergencia entre la estadística, la computación, la minería de datos, el aprendizaje automático y la necesidad de convertir datos en conocimiento útil. Este ensayo analiza el origen y la evolución de la ciencia de datos desde sus antecedentes en el análisis de datos de John W. Tukey hasta su consolidación contemporánea como práctica científica, tecnológica y organizacional. Se revisan fuentes arbitradas de alto impacto para explicar el papel fundador de Tukey, la formalización del término por Cleveland, la importancia del proceso KDD, el contraste entre minería de datos y estadística matemática, el surgimiento de los algoritmos de aprendizaje automático y las sociedades científicas que lideran el campo. El trabajo sostiene que la ciencia de datos no reemplaza a la estadística ni a la informática, sino que las integra en un proceso orientado a la inferencia, la predicción, la explicación, la visualización, la comunicación y la toma de decisiones. Asimismo, se incluye una línea de tiempo como figura y un anexo con código QR para acceder a una versión en línea de dicha trayectoria. La conclusión principal es que la ciencia de datos debe comprenderse como una evolución metodológica y cultural basada en aprender de los datos con rigor, reproducibilidad y responsabilidad.", indent=False)
    para(doc, "Palabras clave: ciencia de datos, minería de datos, aprendizaje automático, estadística, KDD", italic=True, indent=False)
    page_break(doc)
    h1(doc, "Abstract")
    para(doc, "Data science is an interdisciplinary field that emerges from the convergence of statistics, computing, data mining, machine learning, and the need to transform data into useful knowledge. This essay analyzes the origin and evolution of data science from John W. Tukey's work on data analysis to its contemporary consolidation as a scientific, technological, and organizational practice. High-impact peer-reviewed sources are reviewed to explain Tukey's foundational role, Cleveland's formalization of the term, the importance of the KDD process, the contrast between data mining and mathematical statistics, the emergence of machine learning algorithms, and the scientific societies that lead the field. The paper argues that data science does not replace statistics or computer science; rather, it integrates them into a process oriented toward inference, prediction, explanation, visualization, communication, and decision-making. A timeline is included as a figure, and an annex provides a QR code to access an online version of that trajectory. The main conclusion is that data science should be understood as a methodological and cultural evolution based on learning from data with rigor, reproducibility, and responsibility.", indent=False)
    para(doc, "Keywords: data science, data mining, machine learning, statistics, KDD", italic=True, indent=False)
    page_break(doc)


def toc(doc):
    h1(doc, "Tabla de contenido")
    add_updateable_toc(doc)
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(body_section)
    set_page_numbering(body_section, "decimal", 1)
    configure_header(body_section, roman=False)


def body(doc):
    h1(doc, "Origen y Evolución de la Ciencia de Datos")
    h2(doc, "Introducción")
    para(doc, "La ciencia de datos, conocida también como Data Science o DSC, es una de las áreas centrales de la transformación científica, empresarial y social del siglo XXI. Su importancia se debe a que las organizaciones producen volúmenes de datos que ya no pueden interpretarse únicamente con procedimientos manuales o con análisis aislados. Sin embargo, su historia no empieza con el Big Data ni con la inteligencia artificial reciente; sus raíces se encuentran en el análisis estadístico, la computación, el aprendizaje automático, la minería de datos, la visualización y la búsqueda de métodos reproducibles para aprender de la evidencia empírica.")
    para(doc, "Tukey (1962) sostuvo que el análisis de datos debía comprenderse como una actividad intelectual propia, no subordinada por completo a la estadística matemática. Esta idea abrió la puerta a una visión más práctica, exploratoria y computacional del trabajo con datos. Décadas después, Cleveland (2001) propuso expandir las áreas técnicas de la estadística y usó explícitamente el concepto de ciencia de datos para describir una agenda que integraba modelos, computación, pedagogía, colaboración interdisciplinaria y trabajo aplicado.")
    para(doc, "El objetivo de este ensayo es analizar el origen y la evolución de la ciencia de datos, con énfasis en cinco aspectos solicitados por la asignatura: el padre fundador, las primeras revistas o conferencias científicas, la diferencia entre minería de datos y matemáticas estadísticas, el surgimiento de los algoritmos y las sociedades que lideran el campo. Para cumplir ese propósito, se utilizan fuentes arbitradas y de alto impacto, incluyendo trabajos de Tukey, Cleveland, Fayyad, Breiman, Donoho, Dhar, Provost y Fawcett.")

    h2(doc, "Marco de referencia")
    para(doc, "La ciencia de datos puede definirse como un campo interdisciplinario orientado a extraer conocimiento útil, generalizable y comunicable a partir de datos. Donoho (2017) afirma que la historia de la ciencia de datos debe entenderse como una trayectoria de largo plazo que involucra análisis exploratorio, estadística computacional, visualización, investigación reproducible y aprendizaje basado en datos. Desde esta perspectiva, DSC no es simplemente programar en Python o aplicar modelos automáticos; es organizar un proceso completo que inicia con una pregunta y termina con una interpretación útil para la toma de decisiones.")
    para(doc, "Una cita parentética permite sintetizar la naturaleza integradora del campo: la ciencia de datos articula estadística, informática, aprendizaje automático, minería de datos, visualización, evaluación de modelos y conocimiento del dominio (Cleveland, 2001; Dhar, 2013; Donoho, 2017; Provost y Fawcett, 2013). En consecuencia, su evolución debe observarse como una convergencia entre comunidades académicas que antes trabajaban con énfasis diferentes: estadísticos preocupados por la inferencia, informáticos interesados en algoritmos y bases de datos, y científicos aplicados enfocados en resolver problemas reales.")
    para(doc, "El marco de referencia también exige distinguir entre datos, información, conocimiento y decisión. Un conjunto de datos por sí solo no produce valor; requiere limpieza, contexto, modelado, validación y comunicación. Por eso, la ciencia de datos incorpora tareas técnicas y tareas interpretativas. En el plano técnico, usa algoritmos, bases de datos, modelos estadísticos y herramientas computacionales. En el plano epistemológico, pregunta qué se puede saber, con qué nivel de incertidumbre, bajo qué supuestos y con qué consecuencias éticas.")

    h2(doc, "Padre fundador")
    para(doc, "Aunque la ciencia de datos tiene varios antecedentes, John W. Tukey suele reconocerse como el padre intelectual del análisis de datos moderno. En su artículo The Future of Data Analysis, publicado en The Annals of Mathematical Statistics, Tukey (1962) planteó que el análisis de datos era más amplio que la estadística matemática formal. Su contribución fue decisiva porque defendió una práctica centrada en explorar, representar, interpretar y aprender de los datos reales.")
    para(doc, "Donoho (2017) retoma a Tukey como punto de partida de una historia de cincuenta años de ciencia de datos. En esa lectura, Tukey no solo propuso técnicas, sino una cultura científica: usar gráficos, computación, razonamiento exploratorio y juicio empírico para comprender fenómenos. Una cita textual breve resume el núcleo de esa tradición: la ciencia de datos se orienta a “learning from data” (Donoho, 2017, p. 745). Esa expresión no reduce el campo a una frase, pero sí expresa su finalidad principal: aprender de los datos con métodos rigurosos.")
    para(doc, "William S. Cleveland también merece un lugar central. Cleveland (2001) propuso una agenda formal para expandir la estadística hacia la ciencia de datos, incorporando computación, herramientas, teoría, modelos y colaboración con otras áreas. Por tanto, Tukey puede considerarse el fundador conceptual del análisis de datos moderno, mientras que Cleveland aparece como uno de los autores que formalizó el término ciencia de datos dentro de la estadística contemporánea.")

    h2(doc, "Primeras revistas y conferencias científicas")
    para(doc, "La consolidación de la ciencia de datos como campo científico se observa en la aparición de conferencias, revistas especializadas y comunidades de revisión por pares. Una de las trayectorias más influyentes fue Knowledge Discovery in Databases, conocida como KDD. Esta comunidad surgió al conectar inteligencia artificial, bases de datos, estadística y aprendizaje automático. El proceso KDD permitió entender que descubrir conocimiento no equivale únicamente a ejecutar un algoritmo, sino a seleccionar, limpiar, transformar, modelar, evaluar e interpretar datos.")
    para(doc, "Fayyad et al. (1996) organizaron conceptualmente este proceso y definieron la minería de datos como una etapa dentro de una cadena más amplia de descubrimiento de conocimiento. Su artículo en Communications of the ACM sigue siendo una fuente central porque diferencia entre el procedimiento técnico de encontrar patrones y el proceso completo de convertir datos en conocimiento válido. Esta distinción sigue vigente en la ciencia de datos actual, donde el modelado es solo una parte del ciclo de trabajo.")
    para(doc, "Entre las primeras publicaciones importantes se destaca Data Mining and Knowledge Discovery, revista iniciada en 1997 y dedicada a métodos, aplicaciones y fundamentos de la minería de datos. También se reconoce la importancia de Journal of Computational and Graphical Statistics, fundada en 1992, por su aporte a la estadística computacional y gráfica. Más adelante, Journal of Data Science contribuyó a visibilizar el nombre del campo. En conferencias, ACM SIGKDD se convirtió en uno de los espacios académicos más influyentes para investigación en minería de datos, aprendizaje automático aplicado y ciencia de datos.")

    h2(doc, "Minería de datos versus matemáticas estadísticas")
    para(doc, "La minería de datos y las matemáticas estadísticas están relacionadas, pero no son equivalentes. La estadística matemática se apoya en probabilidad, inferencia, estimación, pruebas de hipótesis, intervalos de confianza, modelos y teoría del error. Su preocupación central es determinar qué conclusiones pueden generalizarse desde una muestra hacia una población y con qué incertidumbre. La minería de datos, en cambio, se enfoca en descubrir patrones útiles, novedosos, comprensibles y accionables dentro de grandes conjuntos de datos.")
    para(doc, "Breiman (2001) explicó esta diferencia mediante la idea de dos culturas del modelado estadístico. Una cultura se orienta a construir modelos probabilísticos interpretables sobre el mecanismo que genera los datos; la otra se concentra en producir algoritmos con alta capacidad predictiva. Este contraste es fundamental para entender por qué la ciencia de datos no puede depender solo de una de las dos tradiciones. Si se privilegia únicamente la predicción, se corre el riesgo de producir modelos opacos o asociaciones espurias. Si se privilegia únicamente la teoría formal, se puede perder capacidad para enfrentar datos masivos, heterogéneos y dinámicos.")
    para(doc, "La minería de datos aporta técnicas como clasificación, agrupamiento, reglas de asociación, detección de anomalías y reducción de dimensionalidad. La estadística aporta criterios de validez, control de sesgo, medición de incertidumbre y diseño inferencial. En conjunto, ambas tradiciones permiten que la ciencia de datos descubra patrones y, al mismo tiempo, evalúe si esos patrones son confiables. Por eso, el contraste correcto no es sustitución, sino integración crítica.")

    h2(doc, "Surgimiento de los algoritmos")
    para(doc, "El surgimiento de los algoritmos en ciencia de datos está vinculado con el crecimiento de la computación, las bases de datos y el aprendizaje automático. Durante el siglo XX se desarrollaron métodos como regresión, árboles de decisión, clasificación bayesiana, clustering, redes neuronales y modelos de optimización. Estos métodos adquirieron mayor importancia cuando las organizaciones comenzaron a almacenar datos transaccionales, textos, imágenes, registros digitales y datos de comportamiento a gran escala.")
    para(doc, "Dhar (2013) relaciona la ciencia de datos con la creación de conocimiento accionable y modelos predictivos. Esa orientación predictiva explica por qué los algoritmos se volvieron centrales: permiten aprender patrones desde ejemplos y evaluar su desempeño en datos nuevos. Provost y Fawcett (2013) también resaltan que la ciencia de datos combina principios, procesos y técnicas para analizar datos de manera sistemática y apoyar decisiones basadas en evidencia.")
    para(doc, "Los algoritmos no sustituyen el razonamiento humano. Su valor depende de la calidad de los datos, la formulación del problema, la elección de variables, la validación del modelo y la interpretación de resultados. En ciencia de datos, un modelo puede ser matemáticamente sofisticado y, aun así, ser inadecuado si responde una pregunta mal formulada o si reproduce sesgos. Por esa razón, el surgimiento de los algoritmos debe entenderse junto con la ética, la reproducibilidad, la transparencia y el conocimiento del dominio.")

    h2(doc, "Sociedades que lideran este campo")
    para(doc, "La ciencia de datos se consolidó también por el liderazgo de sociedades científicas. ACM SIGKDD, grupo especializado de la Association for Computing Machinery, ha sido una de las comunidades más influyentes en minería de datos y descubrimiento de conocimiento. Su conferencia KDD reúne investigación académica, aplicaciones industriales y avances metodológicos en aprendizaje automático, análisis de datos y minería de datos (Association for Computing Machinery, 2020).")
    para(doc, "Desde la estadística, la American Statistical Association ha sido relevante por promover la discusión sobre educación estadística, ciencia de datos, inferencia y alfabetización de datos. El Institute of Mathematical Statistics y la International Statistical Institute también han contribuido desde la teoría estadística, la probabilidad y la investigación metodológica. En el ámbito computacional, IEEE y sus sociedades relacionadas con inteligencia computacional han impulsado investigaciones en redes neuronales, aprendizaje automático, sistemas inteligentes y big data.")
    para(doc, "Estas sociedades demuestran que la ciencia de datos no pertenece a una sola disciplina. Su liderazgo está distribuido entre estadística, informática, inteligencia artificial, ingeniería, negocios y ciencias aplicadas. Esa pluralidad es una fortaleza, pero también exige formación rigurosa para evitar trabajos superficiales, modelos sin validación o conclusiones que no respeten la incertidumbre.")

    h2(doc, "Línea de tiempo")
    para(doc, "La Figura 1 resume los hitos principales de la evolución de la ciencia de datos, desde el planteamiento de Tukey hasta la reconstrucción histórica propuesta por Donoho. Esta representación permite observar que DSC surge gradualmente por la interacción entre estadística, computación, minería de datos y aprendizaje automático.", indent=True)
    label = para(doc, "Figura 1", bold=True, indent=False)
    label.paragraph_format.keep_with_next = True
    para(doc, "Línea de tiempo resumida de la evolución de la ciencia de datos", italic=True, indent=False)
    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.add_run().add_picture(str(TIMELINE), width=Inches(6.4))
    para(doc, "Nota. Elaboración propia con base en Tukey (1962), Fayyad et al. (1996), Cleveland (2001), Breiman (2001) y Donoho (2017). El código QR para acceder a una versión en línea se presenta en el Anexo A.", italic=True, indent=False, size=10)

    h2(doc, "Conclusiones")
    para(doc, "La historia de la ciencia de datos muestra una evolución gradual y multidisciplinaria. John W. Tukey aparece como figura fundacional porque defendió el análisis de datos como actividad científica propia, mientras que William S. Cleveland formalizó el término ciencia de datos como una expansión necesaria de la estadística. A partir de los años ochenta y noventa, KDD y la minería de datos aportaron una estructura metodológica para descubrir patrones en grandes bases de datos.")
    para(doc, "El contraste entre minería de datos y matemáticas estadísticas demuestra que ambas perspectivas son necesarias. La estadística proporciona inferencia, medición de incertidumbre y rigor metodológico; la minería de datos y el aprendizaje automático aportan capacidad computacional, descubrimiento de patrones y modelos predictivos. La ciencia de datos surge precisamente de integrar estas tradiciones en un proceso completo de conocimiento.")
    para(doc, "Finalmente, las sociedades científicas y las conferencias especializadas han permitido que el campo se institucionalice y mantenga estándares de revisión académica. Por ello, DSC no debe entenderse como una moda tecnológica, sino como una evolución científica orientada a aprender de los datos, comunicar evidencia y apoyar decisiones responsables en contextos complejos.")
    page_break(doc)


def references(doc):
    h1(doc, "Referencias")
    refs = [
        "Association for Computing Machinery. (2020, 21 de agosto). KDD 2020 showcases brightest minds in data science and AI. https://www.acm.org/media-center/2020/august/sigkdd-2020",
        "Breiman, L. (2001). Statistical modeling: The two cultures. Statistical Science, 16(3), 199–231. https://doi.org/10.1214/ss/1009213726",
        "Cleveland, W. S. (2001). Data science: An action plan for expanding the technical areas of the field of statistics. International Statistical Review, 69(1), 21–26.",
        "Dhar, V. (2013). Data science and prediction. Communications of the ACM, 56(12), 64–73. https://doi.org/10.1145/2500499",
        "Donoho, D. (2017). 50 years of data science. Journal of Computational and Graphical Statistics, 26(4), 745–766. https://doi.org/10.1080/10618600.2017.1384734",
        "Fayyad, U., Piatetsky-Shapiro, G. y Smyth, P. (1996). The KDD process for extracting useful knowledge from volumes of data. Communications of the ACM, 39(11), 27–34. https://doi.org/10.1145/240455.240464",
        "Provost, F. y Fawcett, T. (2013). Data science and its relationship to big data and data-driven decision making. Big Data, 1(1), 51–59. https://doi.org/10.1089/big.2013.1508",
        "Tukey, J. W. (1962). The future of data analysis. The Annals of Mathematical Statistics, 33(1), 1–67. https://doi.org/10.1214/aoms/1177704711",
    ]
    for ref in refs:
        p = para(doc, ref, indent=False)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
    page_break(doc)


def annex(doc, qr_url):
    h1(doc, "Anexo")
    h2(doc, "Anexo A. Código QR de la línea de tiempo")
    para(doc, "El Anexo A presenta el código QR que permite acceder en línea a una versión visual de la línea de tiempo incluida en el cuerpo del documento.", indent=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(QR), width=Inches(1.8))
    para(doc, f"Enlace del QR: {qr_url}", indent=False, size=10)
    h2(doc, "Anexo B. Fuente del logo institucional")
    para(doc, "El logo de la Universidad Autónoma de Santo Domingo usado en la página de presentación procede de Wikimedia Commons, archivo Universidad Autonoma de Santo Domingo.svg, cuya fuente indicada es uasd.edu.do y licencia CC BY-SA 4.0.", indent=False)


def build():
    generate_timeline_image()
    qr_url = generate_qr()
    doc = Document()
    style_doc(doc)
    title_page(doc)
    abstract_pages(doc)
    toc(doc)
    body(doc)
    references(doc)
    annex(doc, qr_url)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
