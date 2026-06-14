from base64 import urlsafe_b64encode
from pathlib import Path

import qrcode
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


GENERAL = Path(__file__).resolve().parents[1]
DOCUMENTS = GENERAL / "Documentos_Academicos"
ASSETS = GENERAL / "Recursos_Visuales"
OUT = DOCUMENTS / "Ensayo_Origen_y_Evolucion_Ciencia_de_Datos_Marlenis_APA7.docx"
LOGO = ASSETS / "uasd_logo.png"
QR = ASSETS / "qr_linea_tiempo_dsc.png"


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def set_keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    p_pr.append(keep)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 2
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_page_number(section.header.paragraphs[0])

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        style.font.bold = True
        style.font.italic = False
        style.paragraph_format.line_spacing = 2
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)


def p(doc, text="", align=None, bold=False, italic=False, indent=True):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 2
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    if indent:
        para.paragraph_format.first_line_indent = Inches(0.5)
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return para


def heading1(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing = 2
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    set_keep_with_next(para)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = True
    return para


def heading2(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.line_spacing = 2
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    set_keep_with_next(para)
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.bold = True
    return para


def page_break(doc):
    doc.add_page_break()


def title_page(doc):
    for _ in range(3):
        p(doc, "", indent=False)

    p(doc, "Origen y Evolución de la Ciencia de Datos", WD_ALIGN_PARAGRAPH.CENTER, bold=True, indent=False)
    p(doc, "Marlenis Judith Concepción Cuevas", WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    p(doc, "Universidad Autónoma de Santo Domingo", WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    p(doc, "INF-8237: Ciencia de Datos I", WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    p(doc, "Maestro Silverio del Orbe", WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    p(doc, "23 de mayo de 2026", WD_ALIGN_PARAGRAPH.CENTER, indent=False)
    page_break(doc)


def abstract_pages(doc):
    heading1(doc, "Resumen")
    p(doc, "La ciencia de datos es un campo interdisciplinario que surge de la convergencia entre el análisis estadístico, la computación, la minería de datos, el aprendizaje automático y la necesidad práctica de extraer conocimiento de grandes volúmenes de información. Este ensayo analiza su origen y evolución a partir de fuentes científicas arbitradas y textos de alto impacto. Se revisa el papel fundacional de John W. Tukey, quien en 1962 propuso una reforma del análisis de datos al defender que el trabajo con datos constituía un campo intelectual propio. También se examina la contribución de William S. Cleveland, quien en 2001 usó explícitamente la expresión ciencia de datos para ampliar las áreas técnicas de la estadística. Se estudian las primeras revistas y conferencias vinculadas al área, especialmente KDD, ACM SIGKDD, Data Mining and Knowledge Discovery y Journal of Data Science. Asimismo, se contrasta la minería de datos con las matemáticas estadísticas, mostrando que la primera enfatiza el descubrimiento de patrones en bases de datos, mientras que la segunda aporta inferencia, probabilidad, estimación y validación de incertidumbre. Finalmente, se describe el surgimiento de los algoritmos de aprendizaje automático y las sociedades científicas que lideran el campo. La conclusión principal es que la ciencia de datos no reemplaza la estadística ni la informática, sino que las integra en una práctica orientada a producir conocimiento reproducible, predictivo, explicativo y útil para la toma de decisiones.", indent=False)
    p(doc, "Palabras clave: ciencia de datos, minería de datos, aprendizaje automático, estadística, KDD", italic=True, indent=False)
    page_break(doc)

    heading1(doc, "Abstract")
    p(doc, "Data science is an interdisciplinary field emerging from the convergence of statistical analysis, computing, data mining, machine learning, and the practical need to extract knowledge from large volumes of information. This essay analyzes its origin and evolution using peer-reviewed scientific sources and high-impact references. It reviews the foundational role of John W. Tukey, who in 1962 called for a reform of data analysis and argued that working with data represented an intellectual field of its own. It also examines William S. Cleveland’s contribution, who explicitly used the term data science in 2001 to expand the technical areas of statistics. The paper discusses early journals and conferences related to the field, particularly KDD, ACM SIGKDD, Data Mining and Knowledge Discovery, and Journal of Data Science. It also contrasts data mining with mathematical statistics, showing that data mining emphasizes pattern discovery in databases, while statistics provides inference, probability, estimation, and uncertainty validation. Finally, the paper describes the emergence of machine learning algorithms and the scientific societies leading the field. The main conclusion is that data science does not replace statistics or computer science; instead, it integrates them into a practice aimed at producing reproducible, predictive, explanatory, and useful knowledge for decision-making.", indent=False)
    p(doc, "Keywords: data science, data mining, machine learning, statistics, KDD", italic=True, indent=False)
    page_break(doc)


def toc(doc):
    heading1(doc, "Tabla de contenido")
    for item in [
        "Origen y Evolución de la Ciencia de Datos",
        "Introducción",
        "Marco de referencia",
        "Padre fundador de la ciencia de datos",
        "Primeras revistas y conferencias científicas",
        "Minería de datos versus matemáticas estadísticas",
        "Surgimiento de los algoritmos",
        "Sociedades que lideran este campo",
        "Línea de tiempo",
        "Conclusiones",
        "Referencias",
        "Anexo",
    ]:
        p(doc, item, indent=False)
    page_break(doc)


def body(doc):
    heading1(doc, "Origen y Evolución de la Ciencia de Datos")

    sections = [
        ("Introducción", [
            "La ciencia de datos, conocida también como Data Science o DSC, se ha convertido en una de las áreas más influyentes del conocimiento contemporáneo porque responde a una realidad concreta: las sociedades, empresas, gobiernos y comunidades científicas producen datos a una velocidad superior a la capacidad humana tradicional de analizarlos. Sin embargo, su origen no se limita al auge reciente del Big Data ni a la popularidad de la inteligencia artificial. La ciencia de datos tiene raíces profundas en la estadística, la computación, la minería de datos, la inteligencia artificial, la visualización y la investigación científica reproducible.",
            "En una cita narrativa, Tukey (1962) defendió que el análisis de datos debía reconocerse como una actividad científica con identidad propia, distinta de la estadística matemática estrictamente formal. Esta idea es central porque anticipa una tensión que todavía existe: la diferencia entre hacer teoría estadística y trabajar con datos reales, incompletos, ruidosos, heterogéneos y dependientes del contexto. Más tarde, Cleveland (2001) propuso ampliar formalmente el campo de la estadística mediante un plan de acción centrado en el analista de datos, la computación, los modelos y la colaboración interdisciplinaria.",
            "El objetivo de este ensayo es investigar la historia y evolución de la ciencia de datos, identificando su padre fundador, sus primeras revistas y conferencias científicas, el contraste entre minería de datos y matemáticas estadísticas, el surgimiento de los algoritmos y las sociedades que lideran el campo. Para ello se utilizan artículos arbitrados, libros académicos y fuentes científicas de alto impacto, siguiendo el estilo APA.",
        ]),
        ("Marco de referencia", [
            "La ciencia de datos puede entenderse como un campo integrador. Donoho (2017) sostiene que la discusión moderna sobre el área debe ubicarse en una tradición más larga de análisis de datos, estadística computacional, investigación reproducible y aprendizaje desde datos. En ese sentido, DSC no es simplemente usar herramientas de programación ni construir modelos automáticos, sino organizar un proceso completo que va desde la formulación del problema hasta la comunicación de resultados.",
            "Una cita parentética ayuda a precisar esta idea: la ciencia de datos se desarrolla en la intersección de estadística, computación, visualización, modelado, inferencia, bases de datos y conocimiento aplicado (Cleveland, 2001; Donoho, 2017; Fayyad et al., 1996). Esta condición interdisciplinaria explica por qué el campo no tiene una sola institución fundadora ni una fecha única de nacimiento, sino una trayectoria evolutiva.",
            "En el plano metodológico, el campo se alimenta de tres líneas principales. La primera proviene de la estadística y su interés por la inferencia, la estimación, el diseño de estudios y la medición de incertidumbre (Casella & Berger, 2002). La segunda proviene de la informática, especialmente de bases de datos, algoritmos, sistemas distribuidos y lenguajes de programación. La tercera proviene del aprendizaje automático y la minería de datos, cuyo foco está en descubrir patrones y producir modelos predictivos útiles.",
        ]),
        ("Padre fundador de la ciencia de datos", [
            "Aunque no existe un consenso absoluto sobre un único padre fundador, la literatura reconoce a John W. Tukey como una figura fundacional del pensamiento que dio origen a la ciencia de datos. En The Future of Data Analysis, Tukey (1962) propuso que el análisis de datos debía ocupar un lugar propio dentro de la actividad científica. Su aporte no consistió en crear un software ni una institución, sino en formular una visión: los datos no debían tratarse únicamente como insumos para probar teorías matemáticas, sino como objetos de exploración, descubrimiento y razonamiento.",
            "Donoho (2017) retoma a Tukey como punto de partida para explicar cincuenta años de ciencia de datos. Según esta lectura, Tukey anticipó una reforma intelectual: aprender de los datos requería métodos, gráficos, experimentación computacional, criterio empírico y comunicación. Una breve cita textual resume su legado conceptual: “learning from data” (Donoho, 2017, p. 745). Esta frase no debe entenderse como una definición completa, sino como una idea organizadora del campo.",
            "No obstante, William S. Cleveland también ocupa un lugar decisivo. Cleveland (2001) usó explícitamente la expresión data science y propuso un plan para expandir las áreas técnicas de la estadística. Por tanto, puede afirmarse que Tukey es el padre intelectual del análisis de datos moderno, mientras Cleveland es uno de los principales formalizadores del término ciencia de datos dentro de la estadística contemporánea.",
        ]),
        ("Primeras revistas y conferencias científicas", [
            "La consolidación de un campo científico suele evidenciarse cuando aparecen revistas, conferencias, sociedades y comunidades de revisión por pares. En ciencia de datos, una de las tradiciones más tempranas fue la de Knowledge Discovery in Databases, conocida como KDD. El primer taller KDD se realizó en 1989, asociado a la comunidad de inteligencia artificial y bases de datos. Luego, la conferencia KDD se convirtió en una de las reuniones científicas más importantes para minería de datos, descubrimiento de conocimiento y análisis avanzado.",
            "Fayyad et al. (1996) fueron fundamentales al organizar conceptualmente el proceso KDD. Estos autores definieron la minería de datos como una etapa dentro de un proceso más amplio que incluye selección, preprocesamiento, transformación, extracción de patrones, evaluación e interpretación. Su artículo en Communications of the ACM es una referencia obligatoria porque ayudó a distinguir entre minería de datos como técnica y descubrimiento de conocimiento como proceso completo.",
            "Entre las primeras revistas relevantes se encuentra Data Mining and Knowledge Discovery, publicada por Springer desde 1997, con Usama Fayyad como editor fundador. Esta revista contribuyó a legitimar científicamente la minería de datos como área de investigación. Otra publicación significativa es Journal of Computational and Graphical Statistics, fundada en 1992, importante para la estadística computacional, visualización y análisis de datos. También se destaca Journal of Data Science, establecido en 2003, cuyo objetivo ha sido publicar investigaciones sobre métodos, computación y aplicaciones de ciencia de datos.",
        ]),
        ("Minería de datos versus matemáticas estadísticas", [
            "La relación entre minería de datos y matemáticas estadísticas es de complementariedad, pero también de contraste. La estadística matemática se fundamenta en probabilidad, inferencia, estimación, pruebas de hipótesis, distribuciones, modelos y teoría del error. Su pregunta central suele ser: ¿qué se puede inferir válidamente sobre una población o un proceso a partir de datos observados? En cambio, la minería de datos pregunta: ¿qué patrones útiles, novedosos o accionables pueden descubrirse en grandes bases de datos?",
            "Breiman (2001) explicó esta diferencia mediante la idea de dos culturas del modelado estadístico. Una cultura se centra en modelos probabilísticos que buscan explicar cómo se generaron los datos; la otra se centra en algoritmos que buscan predecir con precisión la salida a partir de entradas observadas. Este contraste es útil para comprender por qué la minería de datos y el aprendizaje automático no siempre tienen los mismos objetivos que la estadística inferencial clásica.",
            "La minería de datos suele trabajar con grandes volúmenes de datos operacionales, bases transaccionales, registros digitales, textos, redes, imágenes o secuencias. Sus técnicas incluyen clasificación, agrupamiento, reglas de asociación, detección de anomalías, árboles de decisión, redes neuronales y métodos de reducción de dimensionalidad (Han et al., 2011; Witten et al., 2016). La estadística matemática, por su parte, aporta criterios de validez: incertidumbre, sesgo, varianza, significancia, intervalos de confianza, diseño muestral y generalización.",
            "Por eso, no es correcto afirmar que la minería de datos reemplaza a la estadística. Más bien, la ciencia de datos necesita ambas. Sin teoría estadística, los patrones pueden ser espurios; sin minería de datos y computación, muchos fenómenos contemporáneos serían imposibles de explorar a escala. Hastie et al. (2009) integran esta visión al presentar el aprendizaje estadístico como un puente entre inferencia, predicción y minería de datos.",
        ]),
        ("Surgimiento de los algoritmos", [
            "El surgimiento de los algoritmos en ciencia de datos está relacionado con el desarrollo de la inteligencia artificial, el aprendizaje automático, la estadística computacional y el aumento de la capacidad de almacenamiento y procesamiento. Desde mediados del siglo XX se desarrollaron métodos de regresión, clasificación, árboles de decisión, clustering y redes neuronales. Sin embargo, su expansión se aceleró cuando las organizaciones comenzaron a acumular bases de datos masivas y cuando la computación permitió entrenar modelos complejos.",
            "Mitchell (1997) definió el aprendizaje automático como el estudio de programas que mejoran su desempeño en una tarea a partir de la experiencia. Esta definición es importante porque desplaza el foco desde la programación explícita hacia el aprendizaje basado en datos. Posteriormente, Hastie et al. (2009) y James et al. (2021) sistematizaron métodos como regresión lineal, regresión logística, árboles, máquinas de soporte vectorial, métodos de remuestreo, regularización, clustering y aprendizaje supervisado y no supervisado.",
            "El auge de los algoritmos también transformó la noción de evidencia. En la estadística clásica, la evidencia se relaciona con pruebas, estimadores y modelos interpretables. En ciencia de datos, además, se evalúa la capacidad predictiva mediante métricas como exactitud, precisión, sensibilidad, especificidad, error cuadrático medio, validación cruzada y desempeño fuera de muestra. Esta evolución no elimina la inferencia, pero amplía las formas de evaluar el conocimiento generado.",
        ]),
        ("Sociedades que lideran este campo", [
            "Las sociedades científicas han sido esenciales para consolidar la ciencia de datos. Entre las principales se encuentra la Association for Computing Machinery, especialmente a través de ACM SIGKDD, grupo especializado en Knowledge Discovery and Data Mining. Esta sociedad organiza la conferencia KDD, una de las más influyentes del mundo en minería de datos, ciencia de datos aplicada y aprendizaje automático.",
            "También destaca la American Statistical Association, que ha impulsado la integración entre estadística y ciencia de datos mediante publicaciones, secciones técnicas y debates curriculares. El Institute of Mathematical Statistics y la International Statistical Institute han contribuido desde la teoría estadística, la inferencia y la probabilidad. En el área de ingeniería y computación, el IEEE y sus sociedades relacionadas con inteligencia computacional y sistemas de información también lideran investigaciones sobre algoritmos, aprendizaje automático, big data y aplicaciones.",
            "La presencia de estas sociedades confirma que la ciencia de datos es un campo global, institucionalizado y sometido a revisión académica.",
        ]),
    ]

    for title, paragraphs in sections:
        heading2(doc, title)
        for text in paragraphs:
            p(doc, text)

    heading2(doc, "Línea de tiempo")
    p(doc, "Figura 1", italic=True, indent=False)
    p(doc, "Línea de tiempo resumida de la evolución de la ciencia de datos", italic=True, indent=False)
    table = doc.add_table(rows=8, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    rows = [
        ("Año", "Hito", "Importancia"),
        ("1962", "John W. Tukey publica The Future of Data Analysis.", "Plantea el análisis de datos como campo intelectual propio."),
        ("1989", "Primer taller KDD.", "Inicia una comunidad formal sobre descubrimiento de conocimiento en bases de datos."),
        ("1996", "Fayyad, Piatetsky-Shapiro y Smyth sistematizan KDD.", "Distingue minería de datos y proceso de descubrimiento de conocimiento."),
        ("1997", "Inicia Data Mining and Knowledge Discovery.", "Primera revista especializada de alto impacto en minería de datos."),
        ("2001", "Cleveland propone Data Science como ampliación de la estadística.", "Formaliza el término dentro de una agenda técnica."),
        ("2001", "Breiman publica Statistical Modeling: The Two Cultures.", "Contrasta modelado estadístico y modelado algorítmico."),
        ("2017", "Donoho publica 50 Years of Data Science.", "Reconstruye históricamente la identidad del campo."),
    ]
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            if i == 0:
                set_cell_shading(cell, "D9EAF7")
            set_cell_text(cell, value, bold=i == 0)
    p(doc, "Nota. Elaboración propia con base en Tukey (1962), Fayyad et al. (1996), Cleveland (2001), Breiman (2001) y Donoho (2017).", italic=True, indent=False)

    heading2(doc, "Conclusiones")
    p(doc, "La ciencia de datos tiene una evolución compleja y multidisciplinaria. Su origen intelectual puede rastrearse hasta John W. Tukey, quien defendió el análisis de datos como una actividad científica propia. Posteriormente, Cleveland contribuyó a formalizar el término ciencia de datos dentro de una agenda de expansión de la estadística. A partir de los años ochenta y noventa, la minería de datos, KDD y el aprendizaje automático aportaron métodos para descubrir patrones y construir modelos predictivos a gran escala.")
    p(doc, "El contraste entre minería de datos y matemáticas estadísticas muestra que ambas áreas cumplen funciones distintas. La minería de datos enfatiza el descubrimiento, la escala y la utilidad práctica; la estadística aporta inferencia, incertidumbre, rigor metodológico y validación. La ciencia de datos surge precisamente de integrar estas perspectivas, junto con programación, visualización, bases de datos y conocimiento del dominio.")
    p(doc, "Las revistas y conferencias como KDD, Data Mining and Knowledge Discovery, Journal of Computational and Graphical Statistics y Journal of Data Science demuestran que el campo se institucionalizó mediante mecanismos científicos revisados por pares. Asimismo, sociedades como ACM SIGKDD, ASA, IEEE, IMS e ISI continúan liderando su desarrollo. En conclusión, DSC no debe entenderse como una moda tecnológica, sino como una evolución científica orientada a aprender de los datos, producir conocimiento confiable y apoyar decisiones en contextos complejos.")
    page_break(doc)


def references(doc):
    heading1(doc, "Referencias")
    refs = [
        "Breiman, L. (2001). Statistical modeling: The two cultures. Statistical Science, 16(3), 199–231. https://doi.org/10.1214/ss/1009213726",
        "Casella, G., & Berger, R. L. (2002). Statistical inference (2nd ed.). Duxbury.",
        "Cleveland, W. S. (2001). Data science: An action plan for expanding the technical areas of the field of statistics. International Statistical Review, 69(1), 21–26.",
        "Donoho, D. (2017). 50 years of data science. Journal of Computational and Graphical Statistics, 26(4), 745–766. https://doi.org/10.1080/10618600.2017.1384734",
        "Fayyad, U., Piatetsky-Shapiro, G., & Smyth, P. (1996). The KDD process for extracting useful knowledge from volumes of data. Communications of the ACM, 39(11), 27–34. https://doi.org/10.1145/240455.240464",
        "Han, J., Kamber, M., & Pei, J. (2011). Data mining: Concepts and techniques (3rd ed.). Morgan Kaufmann.",
        "Hastie, T., Tibshirani, R., & Friedman, J. (2009). The elements of statistical learning: Data mining, inference, and prediction (2nd ed.). Springer. https://doi.org/10.1007/978-0-387-84858-7",
        "James, G., Witten, D., Hastie, T., & Tibshirani, R. (2021). An introduction to statistical learning: With applications in R (2nd ed.). Springer. https://doi.org/10.1007/978-1-0716-1418-1",
        "Mitchell, T. M. (1997). Machine learning. McGraw-Hill. https://www.cs.cmu.edu/~tom/mlbook.html",
        "Tukey, J. W. (1962). The future of data analysis. The Annals of Mathematical Statistics, 33(1), 1–67.",
        "Witten, I. H., Frank, E., Hall, M. A., & Pal, C. J. (2016). Data mining: Practical machine learning tools and techniques (4th ed.). Morgan Kaufmann.",
    ]
    for ref in refs:
        para = p(doc, ref, indent=False)
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.5)
    page_break(doc)


def annex(doc):
    heading1(doc, "Anexo")
    heading2(doc, "Anexo A. Código QR de la línea de tiempo")
    mermaid = """timeline
    title Evolución de la Ciencia de Datos
    1962 : Tukey propone el futuro del análisis de datos
    1989 : Primer taller KDD
    1996 : Fayyad et al. formalizan el proceso KDD
    1997 : Data Mining and Knowledge Discovery
    2001 : Cleveland propone Data Science
    2001 : Breiman contrasta dos culturas del modelado
    2017 : Donoho reconstruye 50 años de ciencia de datos"""
    url = f"https://mermaid.ink/img/{urlsafe_b64encode(mermaid.encode()).decode()}"
    qrcode.make(url).save(QR)
    qr_p = doc.add_paragraph()
    qr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    qr_p.add_run().add_picture(str(QR), width=Inches(1.8))
    p(doc, f"Enlace del QR: {url}", indent=False)
    heading2(doc, "Anexo B. Fuente del logo institucional")
    p(doc, "El logo de la Universidad Autónoma de Santo Domingo usado en la portada procede de Wikimedia Commons, archivo Universidad Autonoma de Santo Domingo.svg, cuya fuente indicada es uasd.edu.do y licencia CC BY-SA 4.0.", indent=False)


def build():
    doc = Document()
    style_document(doc)
    title_page(doc)
    abstract_pages(doc)
    toc(doc)
    body(doc)
    references(doc)
    annex(doc)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
