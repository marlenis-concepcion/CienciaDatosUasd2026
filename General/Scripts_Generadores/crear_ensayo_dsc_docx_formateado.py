from base64 import urlsafe_b64encode
from pathlib import Path

import qrcode
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


GENERAL = Path(__file__).resolve().parents[1]
DOCUMENTS = GENERAL / "Documentos_Academicos"
ASSETS = GENERAL / "Recursos_Visuales"
OUT = DOCUMENTS / "Ensayo_Origen_y_Evolucion_Ciencia_de_Datos_Marlenis_FORMATEADO_ROMANOS.docx"
LOGO = ASSETS / "uasd_logo.png"
QR = ASSETS / "qr_linea_tiempo_dsc.png"

BLUE = "0073B7"
LIGHT_BLUE = "D9EEF9"
TEAL = "1B9AAA"
GOLD = "F2B705"
PALE_GOLD = "FFF1BF"
GRAY = "F2F4F7"
BLACK = "111827"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=BLACK, size=11, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def set_page_numbering(section, fmt="decimal", start=1):
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))
    pg_num_type.set(qn("w:fmt"), fmt)


def add_page_field(paragraph, roman=False):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE \\* roman" if roman else "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)


def add_colored_rule(doc, fill=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_text(cell, " ", size=2)
    return table


def paragraph(doc, text="", align=None, bold=False, italic=False, color=BLACK, size=12, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return p


def heading(doc, text, level=1, color=BLUE):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(12 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor.from_string(color)
        run.bold = True
    return p


def page_break(doc):
    doc.add_page_break()


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    set_page_numbering(section, "lowerRoman", 1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 14, TEAL),
        ("Heading 3", 12, BLUE),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def apply_section_header_footer(section, roman=False):
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = hp.add_run("INF-8237 Ciencia de Datos I | UASD")
    r.font.name = "Times New Roman"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    r = fp.add_run("Marlenis Judith Concepción Cuevas")
    r.font.name = "Times New Roman"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("555555")
    fp.add_run(" | ")
    add_page_field(fp, roman=roman)


def add_header_footer(doc):
    apply_section_header_footer(doc.sections[0], roman=True)


def add_presentation(doc):
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO), width=Inches(1.45))

    paragraph(doc, "Universidad Autónoma de Santo Domingo (UASD)", WD_ALIGN_PARAGRAPH.CENTER, True, color=BLUE, size=16, space_after=0)
    paragraph(doc, "Facultad de Ciencias", WD_ALIGN_PARAGRAPH.CENTER, True, color=TEAL, size=13, space_after=0)
    paragraph(doc, "Programa de Maestría", WD_ALIGN_PARAGRAPH.CENTER, color=BLACK, size=12, space_after=4)
    add_colored_rule(doc, GOLD)
    paragraph(doc, "Ensayo 1", WD_ALIGN_PARAGRAPH.CENTER, True, color=BLUE, size=18, space_after=0)
    paragraph(doc, "Origen y Evolución de la Ciencia de Datos", WD_ALIGN_PARAGRAPH.CENTER, True, color=BLACK, size=16, space_after=10)

    table = doc.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    rows = [
        ("Asignatura", "INF-8237 Ciencia de Datos I"),
        ("Participante", "Marlenis Judith Concepción Cuevas"),
        ("Maestro", "Silverio del Orbe"),
        ("Horario", "Domingos, 9:00 a.m. a 5:00 p.m."),
        ("Periodo", "17 de mayo al 28 de junio de 2026"),
        ("Duración", "64 horas: 32 teóricas y 32 prácticas"),
    ]
    for i, (label, value) in enumerate(rows):
        set_cell_shading(table.cell(i, 0), LIGHT_BLUE if i % 2 == 0 else PALE_GOLD)
        set_cell_shading(table.cell(i, 1), GRAY if i % 2 == 0 else "FFFFFF")
        set_cell_text(table.cell(i, 0), label, bold=True, color=BLUE)
        set_cell_text(table.cell(i, 1), value)

    paragraph(doc, "Nota académica: El domingo 31 de mayo de 2026 no hay docencia por motivo de la celebración del Día de las Madres.", WD_ALIGN_PARAGRAPH.CENTER, italic=True, color="555555", size=10, space_after=8)
    paragraph(doc, "Santo Domingo, República Dominicana", WD_ALIGN_PARAGRAPH.CENTER, color=BLACK, size=12, space_after=0)
    paragraph(doc, "17 de mayo de 2026", WD_ALIGN_PARAGRAPH.CENTER, color=BLACK, size=12, space_after=0)
    page_break(doc)


def add_abstracts(doc):
    heading(doc, "Resumen", 1)
    paragraph(doc, "La ciencia de datos es un campo interdisciplinario que surge de la convergencia entre el análisis estadístico, la computación, la minería de datos, el aprendizaje automático y la necesidad práctica de extraer conocimiento de grandes volúmenes de información. Este ensayo analiza su origen y evolución a partir de fuentes científicas arbitradas y textos de alto impacto. Se revisa el papel fundacional de John W. Tukey, quien en 1962 propuso una reforma del análisis de datos al defender que el trabajo con datos constituía un campo intelectual propio. También se examina la contribución de William S. Cleveland, quien en 2001 usó explícitamente la expresión ciencia de datos para ampliar las áreas técnicas de la estadística. Se estudian las primeras revistas y conferencias vinculadas al área, especialmente KDD, ACM SIGKDD, Data Mining and Knowledge Discovery y Journal of Data Science. Asimismo, se contrasta la minería de datos con las matemáticas estadísticas, mostrando que la primera enfatiza el descubrimiento de patrones en bases de datos, mientras que la segunda aporta inferencia, probabilidad, estimación y validación de incertidumbre. Finalmente, se describe el surgimiento de los algoritmos de aprendizaje automático y las sociedades científicas que lideran el campo. La conclusión principal es que la ciencia de datos no reemplaza la estadística ni la informática, sino que las integra en una práctica orientada a producir conocimiento reproducible, predictivo, explicativo y útil para la toma de decisiones.")
    paragraph(doc, "Palabras clave: ciencia de datos, minería de datos, aprendizaje automático, estadística, KDD.", bold=True, color=BLUE)
    page_break(doc)

    heading(doc, "Abstract", 1)
    paragraph(doc, "Data science is an interdisciplinary field emerging from the convergence of statistical analysis, computing, data mining, machine learning, and the practical need to extract knowledge from large volumes of information. This essay analyzes its origin and evolution using peer-reviewed scientific sources and high-impact references. It reviews the foundational role of John W. Tukey, who in 1962 called for a reform of data analysis and argued that working with data represented an intellectual field of its own. It also examines William S. Cleveland’s contribution, who explicitly used the term data science in 2001 to expand the technical areas of statistics. The paper discusses early journals and conferences related to the field, particularly KDD, ACM SIGKDD, Data Mining and Knowledge Discovery, and Journal of Data Science. It also contrasts data mining with mathematical statistics, showing that data mining emphasizes pattern discovery in databases, while statistics provides inference, probability, estimation, and uncertainty validation. Finally, the paper describes the emergence of machine learning algorithms and the scientific societies leading the field. The main conclusion is that data science does not replace statistics or computer science; instead, it integrates them into a practice aimed at producing reproducible, predictive, explanatory, and useful knowledge for decision-making.")
    paragraph(doc, "Keywords: data science, data mining, machine learning, statistics, KDD.", bold=True, color=BLUE)
    page_break(doc)


def add_toc(doc):
    heading(doc, "Tabla de contenido", 1)
    items = [
        "1. Origen y Evolución de la Ciencia de Datos",
        "1.1 Introducción",
        "1.2 Marco de referencia",
        "1.3 Padre fundador de la ciencia de datos",
        "1.4 Primeras revistas y conferencias científicas",
        "1.5 Minería de datos versus matemáticas estadísticas",
        "1.6 Surgimiento de los algoritmos",
        "1.7 Sociedades que lideran el campo",
        "1.8 Línea de tiempo",
        "1.9 Conclusiones",
        "Referencias",
        "Anexo",
    ]
    for i, item in enumerate(items):
        p = paragraph(doc, item, color=BLUE if i == 0 else BLACK, bold=i == 0, space_after=2)
        p.paragraph_format.left_indent = Inches(0.2 if "." in item else 0)
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    body_section.top_margin = Inches(1)
    body_section.bottom_margin = Inches(1)
    body_section.left_margin = Inches(1)
    body_section.right_margin = Inches(1)
    set_page_numbering(body_section, "decimal", 1)
    apply_section_header_footer(body_section, roman=False)


def add_body(doc):
    heading(doc, "Origen y Evolución de la Ciencia de Datos", 1)
    add_colored_rule(doc, BLUE)

    sections = [
        ("Introducción", [
            "La ciencia de datos, conocida también como Data Science o DSC, se ha convertido en una de las áreas más influyentes del conocimiento contemporáneo porque responde a una realidad concreta: las sociedades, empresas, gobiernos y comunidades científicas producen datos a una velocidad superior a la capacidad humana tradicional de analizarlos. Sin embargo, su origen no se limita al auge reciente del Big Data ni a la popularidad de la inteligencia artificial. La ciencia de datos tiene raíces profundas en la estadística, la computación, la minería de datos, la inteligencia artificial, la visualización y la investigación científica reproducible.",
            "En una cita narrativa, Tukey (1962) defendió que el análisis de datos debía reconocerse como una actividad científica con identidad propia, distinta de la estadística matemática estrictamente formal. Esta idea es central porque anticipa una tensión que todavía existe: la diferencia entre hacer teoría estadística y trabajar con datos reales, incompletos, ruidosos, heterogéneos y dependientes del contexto. Más tarde, Cleveland (2001) propuso ampliar formalmente el campo de la estadística mediante un plan de acción centrado en el analista de datos, la computación, los modelos y la colaboración interdisciplinaria.",
            "El objetivo de este ensayo es investigar la historia y evolución de la ciencia de datos, identificando su padre fundador, sus primeras revistas y conferencias científicas, el contraste entre minería de datos y matemáticas estadísticas, el surgimiento de los algoritmos y las sociedades que lideran el campo. Para ello se utilizan artículos arbitrados, libros académicos y fuentes científicas de alto impacto, siguiendo el estilo APA.",
        ]),
        ("Marco de referencia", [
            "La ciencia de datos puede entenderse como un campo integrador. Donoho (2017) sostiene que la discusión moderna sobre el área debe ubicarse en una tradición más larga de análisis de datos, estadística computacional, investigación reproducible y aprendizaje desde datos. En ese sentido, DSC no es simplemente usar herramientas de programación ni construir modelos automáticos, sino organizar un proceso completo que va desde la formulación del problema hasta la comunicación de resultados.",
            "Una cita parentética ayuda a precisar esta idea: la ciencia de datos se desarrolla en la intersección de estadística, computación, visualización, modelado, inferencia, bases de datos y conocimiento aplicado (Donoho, 2017; Cleveland, 2001; Fayyad et al., 1996). Esta condición interdisciplinaria explica por qué el campo no tiene una sola institución fundadora ni una fecha única de nacimiento, sino una trayectoria evolutiva.",
            "En el plano metodológico, el campo se alimenta de tres líneas principales. La primera proviene de la estadística y su interés por la inferencia, la estimación, el diseño de estudios y la medición de incertidumbre. La segunda proviene de la informática, especialmente de bases de datos, algoritmos, sistemas distribuidos y lenguajes de programación. La tercera proviene del aprendizaje automático y la minería de datos, cuyo foco está en descubrir patrones y producir modelos predictivos útiles.",
        ]),
        ("Padre fundador de la ciencia de datos", [
            "Aunque no existe un consenso absoluto sobre un único padre fundador, la literatura reconoce a John W. Tukey como una figura fundacional del pensamiento que dio origen a la ciencia de datos. En The Future of Data Analysis, Tukey (1962) propuso que el análisis de datos debía ocupar un lugar propio dentro de la actividad científica. Su aporte no consistió en crear un software ni una institución, sino en formular una visión: los datos no debían tratarse únicamente como insumos para probar teorías matemáticas, sino como objetos de exploración, descubrimiento y razonamiento.",
            "Donoho (2017) retoma a Tukey como punto de partida para explicar cincuenta años de ciencia de datos. Según esta lectura, Tukey anticipó una reforma intelectual: aprender de los datos requería métodos, gráficos, experimentación computacional, criterio empírico y comunicación. Una breve cita textual resume su legado conceptual: “learning from data” (Donoho, 2017, p. 745). Esta frase no debe entenderse como una definición completa, sino como una idea organizadora del campo.",
            "No obstante, William S. Cleveland también ocupa un lugar decisivo. Cleveland (2001) usó explícitamente la expresión data science y propuso un plan para expandir las áreas técnicas de la estadística. Por tanto, puede afirmarse que Tukey es el padre intelectual del análisis de datos moderno, mientras Cleveland es uno de los principales formalizadores del término ciencia de datos dentro de la estadística contemporánea.",
        ]),
        ("Primeras revistas y conferencias científicas", [
            "La consolidación de un campo científico suele evidenciarse cuando aparecen revistas, conferencias, sociedades y comunidades de revisión por pares. En ciencia de datos, una de las tradiciones más tempranas fue la de Knowledge Discovery in Databases, conocida como KDD. El primer taller KDD se realizó en 1989, asociado a la comunidad de inteligencia artificial y bases de datos. Luego, la conferencia KDD se convirtió en una de las reuniones científicas más importantes para minería de datos, descubrimiento de conocimiento y análisis avanzado.",
            "Fayyad et al. (1996) fueron fundamentales al organizar conceptualmente el proceso KDD. Estos autores definieron la minería de datos como una etapa dentro de un proceso más amplio que incluye selección, preprocesamiento, transformación, extracción de patrones, evaluación e interpretación. Su artículo en Communications of the ACM es una referencia obligatoria porque ayudó a distinguir entre minería de datos como técnica y descubrimiento de conocimiento como proceso completo.",
            "Entre las primeras revistas relevantes se encuentra Data Mining and Knowledge Discovery, publicada por Springer desde 1997, con Usama Fayyad como editor fundador. Esta revista contribuyó a legitimar científicamente la minería de datos como área de investigación. Otra publicación significativa es Journal of Computational and Graphical Statistics, fundada en 1992, importante para la estadística computacional, visualización y análisis de datos. También se destaca Journal of Data Science, establecido en 2003, cuyo objetivo ha sido publicar investigaciones sobre métodos, computación y aplicaciones de ciencia de datos.",
        ]),
        ("Minería de datos versus matemáticas estadísticas", [
            "La relación entre minería de datos y matemáticas estadísticas es de complementariedad, pero también de contraste. La estadística matemática se fundamenta en probabilidad, inferencia, estimación, pruebas de hipótesis, distribuciones, modelos y teoría del error. Su pregunta central suele ser: ¿qué se puede inferir válidamente sobre una población o un proceso a partir de datos observados? En cambio, la minería de datos pregunta: ¿qué patrones útiles, novedosos o accionables pueden descubrirse en grandes bases de datos?",
            "Breiman (2001) explicó esta diferencia mediante la idea de dos culturas del modelado estadístico. Una cultura se centra en modelos probabilísticos que buscan explicar cómo se generaron los datos; la otra se centra en algoritmos que buscan predecir con precisión la salida a partir de entradas observadas. Este contraste es útil para comprender por qué la minería de datos y el aprendizaje automático no siempre tienen los mismos objetivos que la estadística inferencial clásica.",
            "Por eso, no es correcto afirmar que la minería de datos reemplaza a la estadística. Más bien, la ciencia de datos necesita ambas. Sin teoría estadística, los patrones pueden ser espurios; sin minería de datos y computación, muchos fenómenos contemporáneos serían imposibles de explorar a escala. Hastie et al. (2009) integran esta visión al presentar el aprendizaje estadístico como un puente entre inferencia, predicción y minería de datos.",
        ]),
        ("Surgimiento de los algoritmos", [
            "El surgimiento de los algoritmos en ciencia de datos está relacionado con el desarrollo de la inteligencia artificial, el aprendizaje automático, la estadística computacional y el aumento de la capacidad de almacenamiento y procesamiento. Desde mediados del siglo XX se desarrollaron métodos de regresión, clasificación, árboles de decisión, clustering y redes neuronales. Sin embargo, su expansión se aceleró cuando las organizaciones comenzaron a acumular bases de datos masivas y cuando la computación permitió entrenar modelos complejos.",
            "Mitchell (1997) definió el aprendizaje automático como el estudio de programas que mejoran su desempeño en una tarea a partir de la experiencia. Esta definición es importante porque desplaza el foco desde la programación explícita hacia el aprendizaje basado en datos. Posteriormente, Hastie et al. (2009) y James et al. (2021) sistematizaron métodos como regresión lineal, regresión logística, árboles, máquinas de soporte vectorial, métodos de remuestreo, regularización, clustering y aprendizaje supervisado y no supervisado.",
            "El auge de los algoritmos también transformó la noción de evidencia. En la estadística clásica, la evidencia se relaciona con pruebas, estimadores y modelos interpretables. En ciencia de datos, además, se evalúa la capacidad predictiva mediante métricas como exactitud, precisión, sensibilidad, especificidad, error cuadrático medio, validación cruzada y desempeño fuera de muestra.",
        ]),
        ("Sociedades que lideran este campo", [
            "Las sociedades científicas han sido esenciales para consolidar la ciencia de datos. Entre las principales se encuentra la Association for Computing Machinery, especialmente a través de ACM SIGKDD, grupo especializado en Knowledge Discovery and Data Mining. Esta sociedad organiza la conferencia KDD, una de las más influyentes del mundo en minería de datos, ciencia de datos aplicada y aprendizaje automático.",
            "También destaca la American Statistical Association, que ha impulsado la integración entre estadística y ciencia de datos mediante publicaciones, secciones técnicas y debates curriculares. El Institute of Mathematical Statistics y la International Statistical Institute han contribuido desde la teoría estadística, la inferencia y la probabilidad. En el área de ingeniería y computación, el IEEE y sus sociedades relacionadas con inteligencia computacional y sistemas de información también lideran investigaciones sobre algoritmos, aprendizaje automático, big data y aplicaciones.",
            "La presencia de estas sociedades confirma que la ciencia de datos es un campo global, institucionalizado y sometido a revisión académica.",
        ]),
    ]

    for title, paras in sections:
        heading(doc, title, 2)
        for text in paras:
            paragraph(doc, text)

    heading(doc, "Línea de tiempo", 2)
    paragraph(doc, "Figura 1", bold=True, color=BLUE, space_after=0)
    paragraph(doc, "Línea de tiempo resumida de la evolución de la ciencia de datos", italic=True, color="555555", space_after=4)
    rows = [
        ("Año", "Hito", "Importancia"),
        ("1962", "John W. Tukey publica The Future of Data Analysis.", "Plantea el análisis de datos como campo intelectual propio."),
        ("1989", "Primer taller KDD.", "Inicia una comunidad formal sobre descubrimiento de conocimiento en bases de datos."),
        ("1996", "Fayyad, Piatetsky-Shapiro y Smyth sistematizan KDD.", "Distingue minería de datos y proceso de descubrimiento de conocimiento."),
        ("1997", "Inicia Data Mining and Knowledge Discovery.", "Primera revista especializada de alto impacto en minería de datos."),
        ("2001", "Cleveland propone Data Science como ampliación de la estadística.", "Formaliza el término dentro de una agenda técnica."),
        ("2001", "Breiman publica Statistical Modeling: The Two Cultures.", "Contrasta modelado estadístico y modelado algorítmico."),
        ("2017", "Donoho publica 50 Years of Data Science.", "Reconstruye históricamente la identidad del campo."),
    ]
    t = doc.add_table(rows=len(rows), cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            cell = t.cell(r, c)
            fill = BLUE if r == 0 else (LIGHT_BLUE if r % 2 else PALE_GOLD)
            set_cell_shading(cell, fill)
            set_cell_text(cell, value, bold=(r == 0 or c == 0), color="FFFFFF" if r == 0 else BLACK, size=10)
    paragraph(doc, "Nota. Elaboración propia con base en Tukey (1962), Fayyad et al. (1996), Cleveland (2001), Breiman (2001) y Donoho (2017).", italic=True, color="555555", size=10)

    heading(doc, "Conclusiones", 2)
    paragraph(doc, "La ciencia de datos tiene una evolución compleja y multidisciplinaria. Su origen intelectual puede rastrearse hasta John W. Tukey, quien defendió el análisis de datos como una actividad científica propia. Posteriormente, Cleveland contribuyó a formalizar el término ciencia de datos dentro de una agenda de expansión de la estadística. A partir de los años ochenta y noventa, la minería de datos, KDD y el aprendizaje automático aportaron métodos para descubrir patrones y construir modelos predictivos a gran escala.")
    paragraph(doc, "El contraste entre minería de datos y matemáticas estadísticas muestra que ambas áreas cumplen funciones distintas. La minería de datos enfatiza el descubrimiento, la escala y la utilidad práctica; la estadística aporta inferencia, incertidumbre, rigor metodológico y validación. La ciencia de datos surge precisamente de integrar estas perspectivas, junto con programación, visualización, bases de datos y conocimiento del dominio.")
    paragraph(doc, "Las revistas y conferencias como KDD, Data Mining and Knowledge Discovery, Journal of Computational and Graphical Statistics y Journal of Data Science demuestran que el campo se institucionalizó mediante mecanismos científicos revisados por pares. Asimismo, sociedades como ACM SIGKDD, ASA, IEEE, IMS e ISI continúan liderando su desarrollo. En conclusión, DSC no debe entenderse como una moda tecnológica, sino como una evolución científica orientada a aprender de los datos, producir conocimiento confiable y apoyar decisiones en contextos complejos.")
    page_break(doc)


def add_references(doc):
    heading(doc, "Referencias", 1)
    refs = [
        "Breiman, L. (2001). Statistical modeling: The two cultures. Statistical Science, 16(3), 199–231. https://doi.org/10.1214/ss/1009213726",
        "Casella, G., & Berger, R. L. (2002). Statistical inference (2nd ed.). Duxbury.",
        "Cleveland, W. S. (2001). Data science: An action plan for expanding the technical areas of the field of statistics. International Statistical Review, 69(1), 21–26.",
        "Donoho, D. (2017). 50 years of data science. Journal of Computational and Graphical Statistics, 26(4), 745–766. https://doi.org/10.1080/10618600.2017.1384734",
        "Fayyad, U., Piatetsky-Shapiro, G., & Smyth, P. (1996). The KDD process for extracting useful knowledge from volumes of data. Communications of the ACM, 39(11), 27–34. https://doi.org/10.1145/240455.240464",
        "Han, J., Kamber, M., & Pei, J. (2011). Data mining: Concepts and techniques (3rd ed.). Morgan Kaufmann.",
        "Hastie, T., Tibshirani, R., & Friedman, J. (2009). The elements of statistical learning: Data mining, inference, and prediction (2nd ed.). Springer. https://doi.org/10.1007/978-0-387-84858-7",
        "James, G., Witten, D., Hastie, T., & Tibshirani, R. (2021). An introduction to statistical learning: With applications in R (2nd ed.). Springer. https://doi.org/10.1007/978-1-0716-1418-1",
        "Mitchell, T. M. (1997). Machine learning. McGraw-Hill. https://www.cs.cmu.edu/~tom/mlbook.html",
        "Tukey, J. W. (1962). The future of data analysis. The Annals of Mathematical Statistics, 33(1), 1–67.",
        "Witten, I. H., Frank, E., Hall, M. A., & Pal, C. J. (2016). Data mining: Practical machine learning tools and techniques (4th ed.). Morgan Kaufmann.",
    ]
    for ref in refs:
        p = paragraph(doc, ref, size=12)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.left_indent = Inches(0.5)
    page_break(doc)


def add_annex(doc):
    heading(doc, "Anexo", 1)
    heading(doc, "Anexo A. Código QR para acceder a la línea de tiempo en línea", 2)
    mermaid = """timeline
    title Evolución de la Ciencia de Datos
    1962 : Tukey propone el futuro del análisis de datos
    1989 : Primer taller KDD
    1996 : Fayyad et al. formalizan el proceso KDD
    1997 : Data Mining and Knowledge Discovery
    2001 : Cleveland propone Data Science
    2001 : Breiman contrasta dos culturas del modelado
    2017 : Donoho reconstruye 50 años de ciencia de datos"""
    encoded = urlsafe_b64encode(mermaid.encode("utf-8")).decode("ascii")
    url = f"https://mermaid.ink/img/{encoded}"
    img = qrcode.make(url)
    img.save(QR)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(QR), width=Inches(1.8))
    paragraph(doc, "El QR dirige a una versión visual en línea de la línea de tiempo elaborada para este ensayo.", WD_ALIGN_PARAGRAPH.CENTER, italic=True, color="555555", size=10)
    paragraph(doc, f"Enlace del QR: {url}", size=9, color=BLUE)
    heading(doc, "Anexo B. Fuente del logo institucional", 2)
    paragraph(doc, "Logo UASD usado en la portada: Wikimedia Commons, archivo Universidad Autonoma de Santo Domingo.svg, fuente indicada: uasd.edu.do, licencia CC BY-SA 4.0.", size=10)


def build():
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    add_presentation(doc)
    add_abstracts(doc)
    add_toc(doc)
    add_body(doc)
    add_references(doc)
    add_annex(doc)
    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()
