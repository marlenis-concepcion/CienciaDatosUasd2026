from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "Articulo_Cientifico_OULAD_APA7.md"
OUTPUT = ROOT / "docs" / "Articulo_Cientifico_OULAD_APA7.docx"


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def add_inline(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part.strip("*`"))
        if part.startswith("**"):
            run.bold = True
        elif part.startswith("*"):
            run.italic = True
        elif part.startswith("`"):
            run.font.name = "Courier New"


def markdown_blocks(text):
    blocks = []
    paragraph_lines = []

    def flush_paragraph():
        if paragraph_lines:
            blocks.append(("paragraph", " ".join(paragraph_lines)))
            paragraph_lines.clear()

    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            flush_paragraph()
            continue
        if raw.endswith("  "):
            flush_paragraph()
            blocks.append(("paragraph", line))
            continue
        if line.startswith("# "):
            flush_paragraph()
            blocks.append(("title", line[2:]))
        elif line.startswith("## "):
            flush_paragraph()
            blocks.append(("heading1", line[3:]))
        elif line.startswith("### "):
            flush_paragraph()
            blocks.append(("heading2", line[4:]))
        elif re.match(r"^\d+\.\s", line):
            flush_paragraph()
            blocks.append(("number", re.sub(r"^\d+\.\s", "", line)))
        elif line.startswith("- "):
            flush_paragraph()
            blocks.append(("bullet", line[2:]))
        else:
            paragraph_lines.append(line)
    flush_paragraph()
    return blocks


def format_paragraph(paragraph, *, indent=True, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph.alignment = alignment
    paragraph.paragraph_format.line_spacing = 2
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Inches(0.5) if indent else Inches(0)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
add_page_number(section.header.paragraphs[0])

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 2
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.first_line_indent = Inches(0.5)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
    style = styles[name]
    style.font.name = "Times New Roman"
    style.font.color.rgb = None
    style.font.bold = True
    style.paragraph_format.line_spacing = 2
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.first_line_indent = Inches(0)

styles["Title"].font.size = Pt(12)
styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
styles["Heading 1"].font.size = Pt(12)
styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
styles["Heading 2"].font.size = Pt(12)
styles["Heading 2"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
styles["Heading 3"].font.size = Pt(12)
styles["Heading 3"].font.italic = True

blocks = markdown_blocks(SOURCE.read_text(encoding="utf-8"))
in_references = False
in_cover = False
in_abstract = False
for kind, text in blocks:
    if kind == "title":
        p = doc.add_paragraph(style="Title")
        add_inline(p, text)
        format_paragraph(p, indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        p.paragraph_format.space_before = Pt(72)
    elif kind == "heading1":
        title = text
        if title in {"Resumen", "Abstract", "Tabla de contenido", "Introducción"}:
            doc.add_page_break()
        p = doc.add_paragraph(title, style="Heading 1")
        format_paragraph(p, indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        in_cover = title == "Presentación"
        in_abstract = title in {"Resumen", "Abstract"}
        in_references = title == "Referencias"
    elif kind == "heading2":
        p = doc.add_paragraph(text, style="Heading 2")
        format_paragraph(p, indent=False)
        in_cover = False
        in_abstract = False
    elif kind == "number":
        p = doc.add_paragraph(style="List Number")
        add_inline(p, text)
        format_paragraph(p, indent=False)
    elif kind == "bullet":
        p = doc.add_paragraph(style="List Bullet")
        add_inline(p, text)
        format_paragraph(p, indent=False)
    else:
        p = doc.add_paragraph()
        add_inline(p, text)
        if in_cover:
            format_paragraph(p, indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        elif in_abstract:
            format_paragraph(p, indent=False)
        else:
            format_paragraph(p)
        if in_references:
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)

doc.save(OUTPUT)
print(OUTPUT.relative_to(ROOT))
