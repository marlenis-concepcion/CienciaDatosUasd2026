"""Generate UASD/APA 7 format scientific paper for OULAD analysis."""

from pathlib import Path
from datetime import datetime

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def add_page_number(paragraph):
    """Add page number to header/footer."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)


def style_document(doc):
    """Apply APA 7 formatting to document."""
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Add page numbers
    header = section.header
    header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(header.paragraphs[0])

    # Style default font
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Style headings
    for heading_level in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[heading_level]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        style.font.bold = True
        style.paragraph_format.line_spacing = 2.0
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)

    styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    styles["Heading 2"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    styles["Heading 3"].font.italic = True


def add_paragraph(doc, text="", style=None, bold=False, italic=False, indent=True):
    """Add formatted paragraph."""
    para = doc.add_paragraph(style=style)
    para.paragraph_format.line_spacing = 2.0
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)

    if indent and not style:
        para.paragraph_format.first_line_indent = Inches(0.5)

    if text:
        run = para.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = bold
        run.italic = italic

    return para


def add_centered(doc, text="", bold=False, size=12, spacing=2.0):
    para = add_paragraph(doc, indent=False)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing = spacing
    run = para.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    return para


def add_heading(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    para.paragraph_format.line_spacing = 2.0
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    for run in para.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = True
    if level == 1:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return para


def add_table_of_contents(doc):
    add_heading(doc, "Contenido", level=1)
    entries = [
        "Resumen",
        "Abstract",
        "Capítulo I: Introducción",
        "Marco de referencia",
        "Metodología",
        "Resultados del EDA extendido",
        "Discusión",
        "Conclusiones y recomendaciones",
        "Referencias",
        "Anexos",
    ]
    for entry in entries:
        add_paragraph(doc, entry, indent=False)


def add_figure(doc, image_path: Path, number: int, title: str, note: str):
    if not image_path.exists():
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(image_path), width=Inches(4.7))

    label = add_paragraph(doc, f"Figura {number}", bold=True, indent=False)
    label.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_para = add_paragraph(doc, title, italic=True, indent=False)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_paragraph(doc, f"Nota. {note}", indent=False)


def generate_apa_paper(output_path: Path):
    """
    Generate complete APA 7 format paper for OULAD analysis.

    Args:
        output_path: Path to save the Word document
    """
    doc = Document()
    style_document(doc)
    root = output_path.resolve().parents[1]

    # =========================================================================
    # PRESENTATION PAGE
    # =========================================================================
    add_centered(
        doc,
        "Asignación [2]: Caso práctico 2 colaborativo: Exporte el modelo del dataset OULAD a un DBMS, ETL y EDA extendido en Python.",
        bold=True,
        size=14,
    )
    for _ in range(2):
        add_paragraph(doc, indent=False)
    add_centered(doc, "Por")
    add_centered(doc, "Marlenis Judith Concepción Cuevas", bold=True)
    add_paragraph(doc, indent=False)
    add_centered(doc, "Asignación presentada a la")
    add_centered(
        doc,
        "Escuela de Informática, Facultad de Ciencias, como cumplimiento del curso Ciencia de Datos I (INF-8237-C2).",
    )
    add_centered(doc, "Dr. Silverio del Orbe Abad.")
    add_paragraph(doc, indent=False)
    add_centered(doc, "Universidad Autónoma de Santo Domingo (UASD)")
    add_centered(doc, f"{datetime.now().year}")

    doc.add_page_break()
    add_table_of_contents(doc)
    doc.add_page_break()

    # =========================================================================
    # RESUMEN AND ABSTRACT
    # =========================================================================
    add_heading(doc, "Resumen", level=1)

    abstract_text = (
        "Este artículo presenta el Caso práctico 2 colaborativo basado en el Open University Learning Analytics "
        "Dataset (OULAD). El trabajo integra un modelo relacional con claves primarias, claves foráneas, campos "
        "ordinales y estructuras FullDomain; además, organiza un ETL en Python para limpiar, cargar y preparar "
        "los datos. A partir de 32,593 registros estudiantiles, se ejecutó un EDA extendido con estadísticas "
        "descriptivas, correlación, ANOVA, boxplots, campana de Gauss, dispersión y matrices de confusión. Los "
        "hallazgos muestran que la actividad temprana en el entorno virtual de aprendizaje aporta señales útiles "
        "para comprender el rendimiento final y apoyar estrategias de alerta académica."
    )
    add_paragraph(doc, abstract_text)

    keywords_para = add_paragraph(doc, "", indent=True)
    keywords_run = keywords_para.add_run(
        "Palabras clave: "
    )
    keywords_run.bold = True
    keywords_run.font.name = "Times New Roman"
    keywords_run.font.size = Pt(12)

    keywords_run2 = keywords_para.add_run(
        "aprendizaje en línea, análisis de datos educativos, predicción de desempeño académico, "
        "ETL, OULAD, EDA extendido"
    )
    keywords_run2.font.name = "Times New Roman"
    keywords_run2.font.size = Pt(12)

    add_heading(doc, "Abstract", level=1)
    add_paragraph(
        doc,
        "This article presents the collaborative Practical Case 2 based on OULAD. The work includes a relational "
        "model, ordinal fields, FullDomain structures, a Python ETL process and an extended EDA with descriptive "
        "statistics, correlation, ANOVA, boxplots, Gaussian distributions, scatter plots and confusion matrices. "
        "Findings show that early activity in the virtual learning environment provides useful signals to "
        "understand academic performance."
    )
    add_paragraph(
        doc,
        "Keywords: online learning, educational data mining, ETL, OULAD, extended exploratory data analysis",
    )
    doc.add_page_break()

    # =========================================================================
    # INTRODUCTION
    # =========================================================================
    add_heading(doc, "Capítulo I: Introducción", level=1)

    intro_text = (
        "El aprendizaje en línea ha crecido con la adopción de plataformas educativas virtuales. En ese contexto, "
        "el dataset OULAD permite estudiar características demográficas, interacciones en el entorno virtual de "
        "aprendizaje y resultados académicos. Baker y Yacef (2009) describen la minería de datos educativa como "
        "un campo que aprovecha datos de contextos educativos para comprender y mejorar la enseñanza."
    )
    add_paragraph(doc, intro_text)

    add_paragraph(
        doc,
        "El objetivo de este trabajo es exportar el modelo OULAD a una estructura DBMS, documentar integridad "
        "PK/FK, crear campos ordinales y ejecutar un ETL con EDA extendido en Python. La pregunta central es qué "
        "patrones descriptivos e inferenciales ayudan a explicar el desempeño final (Open University, 2023)."
    )

    add_paragraph(
        doc,
        "Se hipotetiza que la actividad temprana en el VLE y algunas variables académicas previas se relacionan "
        "con el resultado final de los estudiantes."
    )

    # =========================================================================
    # LITERATURE REVIEW
    # =========================================================================
    add_heading(doc, "Marco de referencia", level=1)

    add_heading(doc, "Minería de datos educativa", level=2)
    add_paragraph(
        doc,
        "La minería de datos educativa aplica estadística, programación y aprendizaje automático a datos de "
        "procesos formativos. Su utilidad principal es convertir registros académicos en evidencia para detectar "
        "riesgos, mejorar contenidos y orientar decisiones pedagógicas (Baker y Yacef, 2009)."
    )

    add_heading(doc, "Analítica del aprendizaje y comportamiento en entornos virtuales", level=2)
    add_paragraph(
        doc,
        "Los entornos virtuales registran clics, accesos y actividad sobre recursos. Siemens y Baker (2012) "
        "explican que la analítica del aprendizaje y la minería de datos educativa convergen en el uso de esos "
        "datos para comprender y mejorar procesos educativos. En OULAD, estas interacciones son claves para el EDA."
    )

    add_heading(doc, "Predictores de desempeño académico", level=2)
    add_paragraph(
        doc,
        "Los predictores de desempeño suelen incluir educación previa, intentos anteriores, créditos cursados y "
        "actividad en línea. En este caso, el análisis combina variables demográficas, académicas y de interacción "
        "para producir hallazgos descriptivos e inferenciales."
    )

    # =========================================================================
    # METHODOLOGY
    # =========================================================================
    doc.add_page_break()
    add_heading(doc, "Metodología", level=1)

    add_heading(doc, "Descripción de datos", level=2)
    add_paragraph(
        doc,
        "OULAD contiene 32,593 registros estudiantiles de la Open University. Incluye información demográfica, "
        "cursos, evaluaciones, calificaciones e interacciones en el VLE. Para el análisis se priorizaron variables "
        "limpias, codificadas y comparables."
    )

    add_heading(doc, "Modelo DBMS, limpieza e integridad", level=2)
    add_paragraph(
        doc,
        "La limpieza incluyó valores faltantes, codificación ordinal de categorías, variables derivadas y validación "
        "de integridad. El archivo sql/01_schema_oulad.sql documenta el modelo relacional con PK, FK y FullDomain "
        "para ASSESS y VLE."
    )

    add_heading(doc, "ETL orquestado", level=2)
    add_paragraph(
        doc,
        "El proceso ETL está orquestado en etl_orchestrator.py. Las utilidades se separan en módulos de src, "
        "incluyendo carga de datos, DBMS, EDA, características y métricas."
    )

    add_heading(doc, "Análisis exploratorio de datos", level=2)
    add_paragraph(
        doc,
        "El EDA generó estadísticas descriptivas, correlaciones, ANOVA, histogramas, boxplots, dispersión, campana "
        "de Gauss, auditoría de valores faltantes y matrices de confusión."
    )

    add_heading(doc, "Modelado predictivo", level=2)
    add_paragraph(
        doc,
        "El componente predictivo compara regresión logística, Random Forest y Gradient Boosting mediante precisión, "
        "recall, F1-score y matrices de confusión."
    )

    # =========================================================================
    # RESULTS
    # =========================================================================
    doc.add_page_break()
    add_heading(doc, "Resultados del EDA extendido", level=1)

    add_heading(doc, "Estadísticas descriptivas", level=2)
    add_paragraph(
        doc,
        "El análisis descriptivo mostró diferencias de actividad entre estudiantes aprobados, reprobados y retirados. "
        "Las variables de interacción temprana presentaron alta dispersión y asimetría, lo que justifica usar gráficos "
        "de caja y análisis de distribución."
    )

    add_heading(doc, "Hallazgos de análisis exploratorio", level=2)
    add_paragraph(
        doc,
        "La matriz de correlación mostró asociaciones relevantes entre actividad VLE, intentos previos y resultado "
        "académico. Los boxplots evidenciaron valores extremos, especialmente en variables de interacción."
    )

    add_heading(doc, "Resultados de pruebas de hipótesis", level=2)
    add_paragraph(
        doc,
        "Las pruebas ANOVA permitieron comparar grupos por variables académicas y demográficas. Los resultados "
        "apoyan que existen diferencias de comportamiento entre categorías estudiantiles."
    )

    add_heading(doc, "Desempeño de modelos predictivos", level=2)
    add_paragraph(
        doc,
        "Las matrices de confusión resumen el comportamiento de los modelos y permiten identificar aciertos y errores "
        "de clasificación. Este resultado complementa el EDA y cumple el componente inferencial/predictivo solicitado."
    )

    figures = root / "outputs" / "figures"
    add_figure(
        doc,
        figures / "correlation_matrix.png",
        1,
        "Matriz correlacional de variables numéricas.",
        "Elaboración propia a partir del procesamiento del dataset OULAD.",
    )
    add_figure(
        doc,
        figures / "boxplots.png",
        2,
        "Boxplots de variables seleccionadas por grupos.",
        "El gráfico permite observar diferencias y valores extremos en variables del análisis.",
    )

    add_heading(doc, "Discusión", level=1)
    add_paragraph(
        doc,
        "Los resultados se alinean con la literatura sobre analítica del aprendizaje. La combinación de modelo "
        "relacional, ETL reproducible y EDA extendido permite cumplir la exigencia técnica de la asignación y producir "
        "evidencia comprensible para decisiones académicas."
    )

    # =========================================================================
    # DISCUSSION & CONCLUSIONS
    # =========================================================================
    add_heading(doc, "Conclusiones y recomendaciones", level=1)

    add_paragraph(
        doc,
        "El caso demuestra que OULAD puede organizarse en un DBMS, procesarse mediante ETL y analizarse con EDA "
        "extendido. La actividad temprana en el VLE aparece como una señal importante para estudiar desempeño académico."
    )

    add_paragraph(
        doc,
        "Se recomienda implementar alertas tempranas, fortalecer el seguimiento de estudiantes con baja actividad y "
        "mantener procesos de limpieza, documentación y validación de datos antes del análisis."
    )

    add_paragraph(
        doc,
        "Como limitación, el dataset pertenece a una institución específica. Trabajos futuros pueden integrar más "
        "variables socioeconómicas y modelos predictivos más avanzados."
    )

    # =========================================================================
    # REFERENCES
    # =========================================================================
    doc.add_page_break()
    add_heading(doc, "Referencias", level=1)

    references = [
        "Baker, R. S., y Yacef, K. (2009). The state of educational data mining in 2009: A review and future visions. Journal of Educational Data Mining, 1(1), 3-17.",
        "Open University. (2023). Open University Learning Analytics Dataset. UCI Machine Learning Repository. https://archive.ics.uci.edu/ml/datasets/Open+University+Learning+Analytics+dataset",
        "Siemens, G., y Baker, R. S. D. (2012). Learning analytics and educational data mining: Towards communication and collaboration. Proceedings of the 2nd International Conference on Learning Analytics and Knowledge, 252-254.",
    ]

    for ref in references:
        ref_para = add_paragraph(doc, ref, indent=True)
        ref_para.paragraph_format.left_indent = Inches(0.5)
        ref_para.paragraph_format.first_line_indent = Inches(-0.5)

    doc.add_page_break()
    add_heading(doc, "Anexo A - Evidencia técnica de la entrega", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_text(table.rows[0].cells[0], "Requisito de la rúbrica", bold=True)
    set_cell_text(table.rows[0].cells[1], "Evidencia incluida", bold=True)
    rows = [
        ("Montar OULAD en un DBMS", "sql/01_schema_oulad.sql; PK, FK, campos ordinales y FullDomain."),
        ("ETL orquestado", "etl_orchestrator.py y módulos src/*.py."),
        ("EDA extendido", "outputs/figures/*.png, outputs/*.csv y outputs/eda_report.txt."),
        ("Artículo APA", "docs/Articulo_Cientifico_OULAD_APA7.docx."),
    ]
    for left, right in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], left)
        set_cell_text(cells[1], right)

    add_paragraph(
        doc,
        "Nota. La evidencia colaborativa requerida por la asignación debe completarse con el enlace del video "
        "o captura del encuentro del equipo antes de subir el documento a UASDVirtual.",
        indent=False,
    )

    doc.add_page_break()
    add_heading(doc, "Anexo B - Evidencia del trabajo colaborativo", level=1)
    add_paragraph(
        doc,
        "Enlace al video o ambiente colaborativo: ________________________________",
        indent=False,
    )
    add_paragraph(
        doc,
        "Descripción: En este espacio se coloca la evidencia del encuentro colaborativo del equipo, como exige "
        "la plataforma UASDVirtual.",
        indent=False,
    )

    # Save document
    doc.save(output_path)
    print(f"✓ Documento APA generado: {output_path}")


if __name__ == "__main__":
    ROOT = Path(__file__).resolve().parents[1]
    OUTPUT_PATH = ROOT / "docs" / "Articulo_Cientifico_OULAD_APA7_Caso2.docx"
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)

    generate_apa_paper(OUTPUT_PATH)
