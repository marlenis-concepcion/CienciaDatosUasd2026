"""Add GitHub link to APA document."""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor

# Document path
ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "docs" / "Articulo_Cientifico_OULAD_APA7.docx"

# GitHub repository URL
GITHUB_URL = "https://github.com/marlenis-concepcion/CienciaDatosUasd2026/tree/main/Unidad_4/Practica_04_Proyecto_Final_OULAD"

def add_github_link():
    """Add GitHub repository link to the document."""
    if not DOC_PATH.exists():
        print(f"❌ Document not found: {DOC_PATH}")
        return False

    try:
        doc = Document(DOC_PATH)

        # Add section for GitHub link at the end
        doc.add_page_break()

        heading = doc.add_heading("Repositorio del Proyecto", level=1)
        heading.paragraph_format.line_spacing = 2.0

        # Add GitHub link
        link_para = doc.add_paragraph()
        link_para.paragraph_format.line_spacing = 2.0
        link_para.paragraph_format.first_line_indent = 0

        link_run = link_para.add_run("Código fuente y documentación disponibles en: ")
        link_run.font.name = "Times New Roman"
        link_run.font.size = Pt(12)

        # Add the URL
        url_run = link_para.add_run(GITHUB_URL)
        url_run.font.name = "Times New Roman"
        url_run.font.size = Pt(12)
        url_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for URL
        url_run.underline = True

        # Save document
        doc.save(DOC_PATH)
        print(f"✅ GitHub link added to document")
        print(f"   Repository: {GITHUB_URL}")
        return True

    except Exception as e:
        print(f"❌ Error adding GitHub link: {e}")
        return False

if __name__ == "__main__":
    success = add_github_link()
    exit(0 if success else 1)
