from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "Articulo_Cientifico_OULAD_APA7.md"
OUTPUT = ROOT / "docs" / "Articulo_Cientifico_OULAD_APA7.docx"


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def add_inline(paragraph, text):
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run(part.strip("*`"))
        if part.startswith("**"):
            run.bold = True
        elif part.startswith("*"):
            run.italic = True
        elif part.startswith("`"):
            run.font.name = "Courier New"


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
add_page_number(section.header.paragraphs[0])

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 2
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.first_line_indent = Inches(0.5)

for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
    style = styles[name]
    style.font.name = "Times New Roman"
    style.font.color.rgb = None
    style.font.bold = True

styles["Title"].font.size = Pt(14)
styles["Title"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
styles["Heading 1"].font.size = Pt(12)
styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
styles["Heading 2"].font.size = Pt(12)
styles["Heading 3"].font.size = Pt(12)
styles["Heading 3"].font.italic = True

lines = SOURCE.read_text(encoding="utf-8").splitlines()
in_references = False
first_heading = True
for raw in lines:
    line = raw.strip()
    if not line:
        continue
    if line.startswith("# "):
        p = doc.add_paragraph(style="Title")
        add_inline(p, line[2:])
        first_heading = False
    elif line.startswith("## "):
        title = line[3:]
        if title in {"Resumen", "Abstract", "Tabla de contenido", "Introducción"}:
            doc.add_page_break()
        p = doc.add_paragraph(title, style="Heading 1")
        in_references = title == "Referencias"
    elif line.startswith("### "):
        doc.add_paragraph(line[4:], style="Heading 2")
    elif re.match(r"^\d+\.\s", line):
        p = doc.add_paragraph(style="List Number")
        add_inline(p, re.sub(r"^\d+\.\s", "", line))
    elif line.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        add_inline(p, line[2:])
    else:
        p = doc.add_paragraph()
        add_inline(p, line)
        if in_references:
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.hanging_indent = Inches(0.5)

doc.save(OUTPUT)
print(OUTPUT.relative_to(ROOT))
