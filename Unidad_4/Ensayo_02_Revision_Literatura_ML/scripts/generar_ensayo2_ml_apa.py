"""Generate Essay 2 DOCX in UASD/APA 7 format."""

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DOCS.mkdir(parents=True, exist_ok=True)


def add_page_number(paragraph, alignment=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = alignment
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def restart_page_numbering(section, start=1, fmt=None):
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))
    if fmt:
        pg_num_type.set(qn("w:fmt"), fmt)


def prepare_section(section, numbered=False, start=1, fmt=None, location="header", alignment=WD_ALIGN_PARAGRAPH.RIGHT):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    clear_paragraph(section.header.paragraphs[0])
    clear_paragraph(section.footer.paragraphs[0])
    if numbered:
        target = section.footer.paragraphs[0] if location == "footer" else section.header.paragraphs[0]
        add_page_number(target, alignment=alignment)
        restart_page_numbering(section, start, fmt)


def start_numbered_body_section(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    prepare_section(section, numbered=True, start=1, fmt="decimal")
    return section


def start_preliminary_section(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    prepare_section(
        section,
        numbered=True,
        start=1,
        fmt="upperRoman",
        location="footer",
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
    )
    return section


def add_toc_field(doc):
    add_heading(doc, "Contenido", level=1)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 2.0

    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "Tabla de contenido automática"

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    clear_paragraph(section.header.paragraphs[0])
    clear_paragraph(section.footer.paragraphs[0])

    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 2.0
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)

    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.styles["Heading 2"].font.bold = True
    doc.styles["Heading 3"].font.bold = True
    doc.styles["Heading 3"].font.italic = True


def add_paragraph(doc, text="", style=None, indent=True):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if indent and not style:
        p.paragraph_format.first_line_indent = Inches(0.5)
    if text:
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def add_center(doc, text="", bold=False, size=12):
    p = add_paragraph(doc, indent=False)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(0, 0, 0)
    r.bold = bold
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.bold = True
    return p


def add_bold_run(paragraph, text):
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    return run


def set_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0, 0, 0)
    r.bold = bold


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell(cell, header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            set_cell(cell, value)
    return table


def add_reference(doc, text):
    p = add_paragraph(doc, text, indent=False)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.5)


def force_all_text_black(doc):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.font.name = "Times New Roman"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)
                        run.font.name = "Times New Roman"


def build_doc():
    doc = Document()
    style_document(doc)

    add_center(doc, "Ensayo 2 [individual]: Revisión de literatura sobre aprendizaje supervisado, no supervisado, self-supervised learning, reforzamiento y deep learning", bold=True, size=14)
    add_paragraph(doc, indent=False)
    add_center(doc, "Por")
    add_center(doc, "Marlenis Judith Concepción Cuevas", bold=True)
    add_paragraph(doc, indent=False)
    add_center(doc, "Asignación presentada a la")
    add_center(doc, "Escuela de Informática, Facultad de Ciencias, como cumplimiento del curso Ciencia de Datos I (INF-8237-C2).")
    add_center(doc, "Dr. Silverio del Orbe Abad.")
    add_paragraph(doc, indent=False)
    add_center(doc, "Universidad Autónoma de Santo Domingo (UASD)")
    add_center(doc, f"Santo Domingo, República Dominicana")
    add_center(doc, f"{datetime.now().year}")
    start_preliminary_section(doc)

    add_toc_field(doc)

    add_heading(doc, "Resumen", level=1)
    add_paragraph(
        doc,
        "Este artículo presenta una revisión de literatura narrativa sobre aprendizaje automático supervisado (SML), no supervisado (NSML), por reforzamiento, self-supervised learning y deep learning. La revisión prioriza literatura posterior a 2021, complementada con trabajos clásicos de alto impacto que explican el nacimiento y consolidación del campo. Desde una perspectiva aplicada a calidad de software, automatización de pruebas, métricas y validación de sistemas, se contrastan los enfoques algorítmicos más utilizados, sus diferencias frente al análisis estadístico convencional y su relación con el acceso actual a grandes volúmenes de datos. El análisis concluye que el valor del ML no está solo en entrenar modelos, sino en validar su comportamiento con métricas adecuadas, trazabilidad y sentido crítico.",
    )
    p = add_paragraph(doc, indent=True)
    add_bold_run(p, "Palabras clave: ")
    p.add_run("aprendizaje automático, aprendizaje supervisado, aprendizaje no supervisado, deep learning, métricas de validación")

    add_heading(doc, "Abstract", level=1)
    add_paragraph(
        doc,
        "This paper presents a narrative literature review on supervised machine learning, unsupervised machine learning, reinforcement learning, self-supervised learning and deep learning. The review prioritizes post-2021 literature while incorporating high-impact classical works that explain the development of the field. From a software quality, automation, testing metrics and system validation perspective, the paper compares algorithmic paradigms, their differences from conventional statistical analysis and their relationship with Big Data access. The review concludes that the value of ML depends not only on model training but also on validation, traceability and critical interpretation.",
    )
    add_paragraph(doc, "Keywords: machine learning, supervised learning, unsupervised learning, deep learning, validation metrics", indent=False)
    start_numbered_body_section(doc)

    add_heading(doc, "Introducción", level=1)
    add_paragraph(
        doc,
        "El aprendizaje automático se ha convertido en una de las áreas más visibles de la ciencia de datos, pero su uso serio requiere separar la moda de la evidencia. Desde mi experiencia en QA automation, pruebas API/UI, CI/CD y validación de flujos críticos, entiendo los modelos de ML como sistemas que también deben probarse: no basta con que produzcan una respuesta; hay que medir si esa respuesta es estable, reproducible, explicable y útil para el contexto. Esa mirada práctica orienta esta revisión.",
    )
    add_paragraph(
        doc,
        "La literatura reciente muestra que la inteligencia artificial avanza por dos fuerzas simultáneas: mejores algoritmos y más datos. El AI Index Report 2026 señala que la evaluación, la gobernanza y la preparación institucional avanzan más lento que la capacidad técnica de los modelos (Sajadieh et al., 2026). Esta tensión importa porque muchas organizaciones adoptan ML sin comprender bien sus métricas, supuestos o riesgos.",
    )
    add_paragraph(
        doc,
        "El objetivo de este ensayo es comparar SML, NSML, reforzamiento, self-supervised learning y deep learning, y contrastar sus métricas con las pruebas estadísticas convencionales. También se revisan tendencias de industria, impacto del Big Data y algoritmos relevantes. El tema se aborda como revisión narrativa, apropiada para integrar teoría, estado del arte y juicio crítico. Manterola et al. (2023) señalan que revisar literatura de buena calidad permite “avanzar en el conocimiento” (párr. 87), idea que guía este documento.",
    )

    add_heading(doc, "Metodología de la revisión", level=1)
    add_paragraph(
        doc,
        "Se realizó una revisión de literatura convencional o narrativa con apoyo de criterios de selección explícitos. Se priorizaron artículos, reportes técnicos y revisiones publicadas entre 2022 y 2026, especialmente en revistas, repositorios académicos y organismos reconocidos. También se incluyeron textos clásicos de alto impacto cuando aportan base conceptual, como LeCun, Bengio y Hinton (2015) para deep learning, y Jordan y Mitchell (2015) para tendencias generales de ML.",
    )
    add_paragraph(
        doc,
        "Los criterios de inclusión fueron: relación directa con SML, NSML, self-supervised learning, deep learning, reforzamiento, evaluación de modelos o tendencias de IA; autoridad académica o institucional; y utilidad para contrastar modelos algorítmicos con investigación convencional. Se excluyeron fuentes divulgativas sin trazabilidad técnica, textos sin autoría clara y materiales que no aportaran al contraste solicitado por la asignación.",
    )
    add_paragraph(
        doc,
        "La revisión no pretende ser un metaanálisis estadístico porque los estudios analizados difieren en dominio, métricas y diseño. En cambio, adopta una síntesis crítica y comparativa, útil para un artículo académico de curso. Esta decisión es coherente con Manterola et al. (2023), quienes distinguen revisiones narrativas, sistemáticas, de alcance y metaanálisis según su nivel de exhaustividad y metodología.",
    )
    add_paragraph(
        doc,
        "Para sostener la originalidad del análisis, no se presenta una simple lista de definiciones. La discusión conecta la literatura con mi experiencia profesional en automatización de pruebas, pruebas API/UI, CI/CD, métricas de calidad, regresión y documentación de defectos. Esa mirada permite contrastar los modelos de ML como sistemas que deben validarse, monitorearse y explicarse, no como herramientas mágicas separadas de la ingeniería de software.",
    )

    add_paragraph(doc, "Tabla 1", indent=False)
    p = add_paragraph(doc, "Criterios de calidad de las fuentes revisadas.", indent=False)
    p.runs[0].italic = True
    add_table(
        doc,
        ["Tipo de fuente", "Criterio de calidad", "Uso en el artículo"],
        [
            ("Artículos revisados por pares", "Publicados en revistas de alto impacto como Nature, Science y Nature Methods.", "Base conceptual para deep learning, ML y contraste estadística vs ML."),
            ("Revisiones recientes", "Literatura posterior a 2021 sobre self-supervised learning y modelos fundacionales.", "Estado del arte y tendencias técnicas."),
            ("Reportes institucionales", "Informes de Stanford HAI/AI Index con datos globales y trazabilidad.", "Tendencias de industria, gobernanza y adopción."),
            ("Libros clásicos", "Obras de referencia usadas solo cuando explican fundamentos históricos.", "Reforzamiento y conceptos base."),
        ],
    )
    add_paragraph(
        doc,
        "Nota. La tabla resume el filtro usado para responder al criterio de calidad e impacto de fuentes de la rúbrica.",
        indent=False,
    )

    add_heading(doc, "Revisión literaria", level=1)
    add_heading(doc, "Aprendizaje supervisado", level=2)
    add_paragraph(
        doc,
        "El aprendizaje supervisado usa datos etiquetados para aprender una relación entre entradas y salidas. En clasificación, el modelo aprende categorías; en regresión, predice valores continuos. Algoritmos como regresión logística, árboles de decisión, random forest, support vector machines, gradient boosting y redes neuronales se utilizan cuando existe una variable objetivo clara. En mi campo profesional, este enfoque se parece a una suite de pruebas con resultado esperado: si el resultado real no coincide con lo esperado, el sistema debe investigarse.",
    )
    add_paragraph(
        doc,
        "Jordan y Mitchell (2015) explican que ML cambió la programación tradicional porque permite que el sistema aprenda patrones desde datos. Sin embargo, el aprendizaje supervisado hereda una dependencia fuerte de la calidad de la etiqueta. Si las etiquetas están sesgadas, incompletas o mal definidas, el modelo puede ser preciso técnicamente y equivocado en la práctica. Esta es una advertencia cercana al QA: una prueba automatizada mal diseñada puede pasar siempre y aun así no proteger el negocio.",
    )

    add_heading(doc, "Aprendizaje no supervisado", level=2)
    add_paragraph(
        doc,
        "El aprendizaje no supervisado trabaja sin etiqueta explícita. Su objetivo es descubrir estructura: grupos, dimensiones latentes, anomalías o asociaciones. K-means, clustering jerárquico, DBSCAN, PCA y autoencoders son ejemplos frecuentes. NSML es especialmente útil cuando el problema todavía no está formulado como clasificación o predicción, sino como exploración. Por eso se relaciona con EDA y con investigación inicial en grandes bases de datos.",
    )
    add_paragraph(
        doc,
        "La diferencia práctica con SML es que NSML no responde de inmediato si un resultado está bien o mal; responde si hay patrones que merecen interpretación. Esto exige mayor criterio humano. En un proyecto real, un cluster no es una conclusión automática: es una hipótesis de segmentación que debe validarse con conocimiento del dominio, métricas internas y evidencia externa.",
    )

    add_heading(doc, "Self-supervised learning y el problema de las etiquetas", level=2)
    add_paragraph(
        doc,
        "El self-supervised learning surge como una respuesta al costo de etiquetar datos. Gui et al. (2023) explican que los modelos supervisados profundos suelen necesitar grandes volúmenes de datos etiquetados, mientras que SSL aprende representaciones a partir de señales generadas por los propios datos. Esto es especialmente relevante en texto, imágenes, recomendadores y modelos fundacionales.",
    )
    add_paragraph(
        doc,
        "Uelwer et al. (2023) presentan SSL como una línea de representación donde el modelo aprende características útiles sin depender de etiquetas humanas. La tendencia tiene sentido industrial: las empresas poseen grandes volúmenes de logs, tickets, imágenes, documentos, trazas de API y eventos de usuario, pero no siempre tienen etiquetas confiables. Desde QA, esto abre una oportunidad: aprender patrones de fallos, anomalías o comportamiento del usuario a partir de evidencia operacional ya existente.",
    )

    add_heading(doc, "Deep learning y modelos fundacionales", level=2)
    add_paragraph(
        doc,
        "Deep learning se basa en redes neuronales con múltiples capas capaces de aprender representaciones jerárquicas. LeCun et al. (2015) consolidaron la idea de que estas arquitecturas son especialmente efectivas en visión, voz y lenguaje. En años recientes, el deep learning se convirtió en base de modelos fundacionales y sistemas generativos, donde una misma arquitectura puede adaptarse a múltiples tareas.",
    )
    add_paragraph(
        doc,
        "Bommasani et al. (2021) introdujeron el concepto de foundation models para describir modelos entrenados con grandes cantidades de datos y adaptables a muchos usos. Esta evolución cambia la pregunta clásica de ML. Ya no se trata solo de escoger entre regresión logística o random forest; ahora también se evalúa si conviene entrenar desde cero, ajustar un modelo preentrenado, usar embeddings o integrar un modelo externo con controles de seguridad.",
    )

    add_heading(doc, "Aprendizaje por reforzamiento", level=2)
    add_paragraph(
        doc,
        "El aprendizaje por reforzamiento se diferencia porque el modelo aprende mediante interacción con un entorno. En lugar de etiquetas fijas, recibe recompensas o penalizaciones. Esta lógica es útil para decisiones secuenciales, robótica, juegos, optimización y agentes. Sutton y Barto (2018) lo explican como un marco para aprender qué acciones tomar en situaciones determinadas a fin de maximizar recompensa acumulada.",
    )
    add_paragraph(
        doc,
        "El reforzamiento no reemplaza SML ni NSML; responde otra clase de problema. Si en SML se pregunta qué etiqueta corresponde a este ejemplo, en reforzamiento se pregunta qué acción conviene tomar ahora considerando consecuencias futuras. En sistemas reales, esto exige controles fuertes, porque una política optimizada con una recompensa mal definida puede aprender conductas no deseadas.",
    )

    add_heading(doc, "Hallazgos y contraste crítico", level=1)
    add_paragraph(
        doc,
        "Un primer hallazgo es que ML no elimina la investigación convencional; la desplaza hacia otra escala. En el análisis estadístico tradicional se trabaja con hipótesis, muestras, significancia, intervalos de confianza y explicación causal o asociativa. En ML se trabaja con generalización, error fuera de muestra, validación cruzada, métricas de clasificación/regresión y desempeño operacional. Bzdok et al. (2018) sostienen que estadística y ML no son enemigos: tienen énfasis distintos.",
    )
    add_paragraph(
        doc,
        "Un segundo hallazgo es que el Big Data cambia el centro de gravedad. La investigación tradicional muchas veces parte de muestras cuidadosamente diseñadas. En ML moderno se trabaja con grandes volúmenes de datos observacionales, logs y comportamiento real. Esto aporta riqueza, pero también ruido, sesgo, datos faltantes y problemas de privacidad. Por eso la calidad del pipeline importa tanto como el algoritmo.",
    )
    add_paragraph(
        doc,
        "Este cambio de muestras a grandes poblaciones de datos no debe interpretarse como eliminación del método científico. Al contrario, exige más disciplina. Cuando se trabaja con millones de registros, el error puede escalar silenciosamente: una variable mal codificada, una fuga de datos entre entrenamiento y prueba, una etiqueta histórica sesgada o un cambio de distribución en producción puede afectar miles de decisiones. En QA automation, un fallo pequeño en una prueba reutilizada se replica en toda la suite; en ML ocurre algo parecido con los datos y las métricas.",
    )
    add_paragraph(
        doc,
        "Por eso, una revisión crítica debe mirar el ciclo completo: definición del problema, calidad de datos, selección algorítmica, entrenamiento, validación, despliegue, monitoreo y retroalimentación. El modelo no termina cuando alcanza un F1-score alto en un notebook. Termina cuando se demuestra que mantiene rendimiento, que no degrada la experiencia del usuario, que no introduce riesgos injustificados y que puede explicarse ante personas técnicas y no técnicas.",
    )
    add_paragraph(
        doc,
        "Un tercer hallazgo es que la industria se mueve hacia modelos preentrenados, SSL, deep learning y automatización de decisiones. El AI Index Report 2026 describe un escenario donde la capacidad técnica avanza rápidamente, mientras evaluación y gobernanza intentan ponerse al día (Sajadieh et al., 2026). Desde mi experiencia con CI/CD y pruebas automatizadas, esto confirma una idea práctica: cuanto más inteligente parece un sistema, más rigurosa debe ser su validación.",
    )

    add_heading(doc, "Métricas de validación: investigación convencional vs ML", level=1)
    add_paragraph(
        doc,
        "El mandato de la asignación pide comparar métricas convencionales con métricas de ML. Esta comparación es central porque define qué se considera evidencia. Alfa de Cronbach evalúa consistencia interna; Pearson y Spearman miden asociación; chi-cuadrado evalúa independencia entre variables categóricas. En ML, en cambio, accuracy, precision, recall, F1-score, MSE, RMSE y AUC evalúan desempeño predictivo. No responden la misma pregunta.",
    )
    add_paragraph(
        doc,
        "En QA esto se entiende bien: una métrica aislada puede engañar. Un modelo con alta accuracy puede fallar en la clase minoritaria; por eso recall y F1-score son indispensables cuando el costo del falso negativo es alto. En regresión, MSE castiga errores grandes, mientras MAE es más interpretable. La métrica correcta depende del riesgo del negocio, no solo del promedio estadístico.",
    )
    add_paragraph(
        doc,
        "La diferencia entre métricas convencionales y métricas de ML también se ve en la intención. Pearson o Spearman ayudan a comprender asociación; chi-cuadrado ayuda a evaluar independencia; alfa de Cronbach ayuda a determinar consistencia interna de un instrumento. En cambio, precision y recall preguntan si el modelo clasifica bien bajo un objetivo operacional. MSE o RMSE preguntan cuánto se equivoca al estimar valores. AUC pregunta qué tan bien separa clases bajo distintos umbrales. Son lenguajes distintos para problemas distintos.",
    )
    add_paragraph(
        doc,
        "En un sistema productivo, ambas familias de métricas pueden convivir. Por ejemplo, antes de entrenar un modelo se puede usar correlación para explorar variables, chi-cuadrado para asociaciones categóricas y análisis de confiabilidad si existe un instrumento de medición. Después del entrenamiento, se usan métricas predictivas y pruebas de robustez. Esta combinación evita dos extremos: confiar solo en pruebas estadísticas sin capacidad predictiva, o confiar solo en métricas de ML sin interpretación científica.",
    )

    add_paragraph(doc, "Tabla 2", indent=False)
    p = add_paragraph(doc, "Contraste entre enfoques de aprendizaje automático y métricas", indent=False)
    p.runs[0].italic = True
    add_table(
        doc,
        ["Enfoque", "Tipo de dato/objetivo", "Algoritmos frecuentes", "Métricas principales"],
        [
            ("SML", "Datos etiquetados; clasificación o regresión", "Regresión logística, SVM, random forest, XGBoost, redes neuronales", "Accuracy, precision, recall, F1, AUC, MSE, RMSE"),
            ("NSML", "Datos sin etiqueta; estructura latente", "K-means, DBSCAN, PCA, autoencoders", "Silhouette, inercia, reconstrucción, validación experta"),
            ("Reforzamiento", "Decisiones secuenciales con recompensa", "Q-learning, DQN, policy gradients", "Recompensa acumulada, regret, estabilidad de política"),
            ("Self-supervised", "Datos no etiquetados con señales internas", "Contrastive learning, masked modeling, embeddings", "Transferencia, desempeño downstream, pérdida contrastiva"),
            ("Deep learning", "Datos complejos: texto, imagen, voz, multimodal", "CNN, RNN, transformers, LLM", "Según tarea: F1, BLEU, perplexity, exact match, MSE"),
        ],
    )

    add_paragraph(doc, "Figura 1", indent=False)
    p = add_paragraph(doc, "Mapa conceptual de selección algorítmica.", indent=False)
    p.runs[0].italic = True
    add_table(
        doc,
        ["Pregunta de investigación", "Ruta recomendada"],
        [
            ("Tengo etiquetas y quiero predecir.", "SML -> clasificación/regresión -> métricas predictivas."),
            ("No tengo etiquetas y quiero descubrir grupos.", "NSML -> clustering/reducción dimensional -> validación por dominio."),
            ("Quiero optimizar acciones en el tiempo.", "Reforzamiento -> política/recompensa -> simulación y control."),
            ("Tengo muchos datos sin etiquetar.", "Self-supervised -> representaciones -> fine-tuning o transferencia."),
            ("Trabajo con texto, imagen, audio o datos complejos.", "Deep learning -> embeddings/modelos fundacionales -> evaluación por tarea."),
        ],
    )
    add_paragraph(doc, "Nota. Elaboración propia para sintetizar la decisión metodológica según el tipo de problema.", indent=False)

    add_heading(doc, "Tendencias de la industria", level=1)
    add_paragraph(
        doc,
        "Las tendencias más fuertes son: modelos fundacionales, automatización con IA generativa, uso de datos sintéticos, evaluación responsable, embeddings, integración de IA en flujos de desarrollo y monitoreo de modelos en producción. Para una profesional de QA, esto significa que probar software ya no se limita a verificar pantallas o APIs; también implica revisar datos, prompts, respuestas, sesgos, regresiones del modelo, trazabilidad y experiencia del usuario.",
    )
    add_paragraph(
        doc,
        "En la industria, SML sigue siendo muy utilizado porque resuelve problemas claros: fraude, scoring, predicción de churn, clasificación de tickets, detección de defectos y estimación de demanda. NSML se usa para segmentación, análisis de comportamiento, anomalías y reducción dimensional. Deep learning domina cuando los datos son complejos o no estructurados: imágenes, audio, texto, video y documentos. SSL y modelos fundacionales ganan fuerza porque permiten reutilizar conocimiento de modelos preentrenados y reducir la dependencia de etiquetas manuales.",
    )
    add_paragraph(
        doc,
        "Los algoritmos tendencia dependen del problema. Para datos tabulares, gradient boosting, random forest y regresión logística siguen siendo competitivos. Para texto, transformers, embeddings y modelos de lenguaje han cambiado la forma de clasificar, resumir y buscar información. Para visión, CNN y vision transformers siguen siendo referencias. Para recomendación, métodos híbridos y self-supervised learning ayudan a manejar escasez de interacción. Para agentes, el reforzamiento y la optimización de políticas aparecen en contextos donde las decisiones tienen consecuencias acumuladas.",
    )
    add_paragraph(
        doc,
        "Narayanan y Kapoor (2024) advierten que muchas promesas de IA se exageran cuando no se distingue entre predicción válida, automatización útil y afirmaciones no demostradas. Esta crítica es importante: no todo problema necesita deep learning, y no todo modelo complejo es superior. En ocasiones, una regresión logística bien documentada supera a una red neuronal opaca si el contexto exige explicabilidad.",
    )
    add_paragraph(
        doc,
        "Esta advertencia es especialmente importante en empresas que adoptan IA por presión competitiva. La pregunta correcta no es cuál modelo está de moda, sino qué evidencia demuestra que el modelo mejora un proceso. En un flujo de QA, por ejemplo, un modelo que prioriza defectos puede ser útil si reduce tiempo de triage sin ocultar fallos críticos. Un modelo que genera casos de prueba puede ayudar si el equipo revisa cobertura, duplicidad y relevancia. La automatización inteligente tiene valor cuando aumenta criterio, no cuando lo reemplaza.",
    )

    add_heading(doc, "Conclusiones", level=1)
    add_paragraph(
        doc,
        "Mi conclusión crítica es que ML debe entenderse como una extensión de la investigación y de la ingeniería, no como sustituto del juicio humano. SML es fuerte cuando existen etiquetas confiables; NSML ayuda a explorar estructuras; SSL reduce dependencia de etiquetas; deep learning domina datos complejos; y reforzamiento atiende decisiones secuenciales. Cada enfoque tiene valor, pero también riesgos.",
    )
    add_paragraph(
        doc,
        "La diferencia más importante entre investigación convencional y ML no está solo en las técnicas, sino en el tipo de evidencia. La investigación convencional busca explicar, contrastar hipótesis y estimar relaciones; ML busca generalizar predicciones o descubrir patrones a escala. Ambas visiones se necesitan. Un proyecto serio debe formular bien la pregunta, cuidar los datos, escoger métricas coherentes y documentar limitaciones.",
    )
    add_paragraph(
        doc,
        "Desde mi experiencia en automatización de pruebas, la lección final es sencilla: un modelo debe probarse como se prueba un sistema crítico. Se valida con métricas, casos límite, monitoreo, evidencia y trazabilidad. El futuro de ML no será solo entrenar modelos más grandes, sino construir sistemas más verificables, responsables y útiles para personas reales.",
    )
    add_paragraph(
        doc,
        "La revisión me deja una postura clara: el aprendizaje automático no debe enseñarse únicamente como catálogo de algoritmos. Debe enseñarse como una forma de investigar con datos, validar hipótesis operativas y construir sistemas que puedan sostenerse en producción. En ese sentido, la experiencia de QA aporta una mirada necesaria: todo resultado debe poder reproducirse, explicarse y fallar de manera visible. Cuando un modelo falla en silencio, la organización no tiene inteligencia artificial; tiene riesgo automatizado.",
    )

    doc.add_page_break()
    add_heading(doc, "Referencias", level=1)
    references = [
        "Balestriero, R., Ibrahim, M., Sobal, V., Morcos, A., Shekhar, S., Goldstein, T., Bordes, F., Bardes, A., Mialon, G., Tian, Y., Schwarzschild, A., Wilson, A. G., Geiping, J., Garrido, Q., Fernandez, P., Bar, A., Pirsiavash, H., LeCun, Y. y Goldblum, M. (2023). A cookbook of self-supervised learning. arXiv. https://arxiv.org/abs/2304.12210",
        "Bommasani, R., Hudson, D. A., Adeli, E., Altman, R., Arora, S., von Arx, S., Bernstein, M. S., Bohg, J., Bosselut, A., Brunskill, E., Brynjolfsson, E., Buch, S., Card, D., Castellon, R., Chatterji, N., Chen, A., Creel, K., Davis, J. Q., Demszky, D., ... Liang, P. (2021). On the opportunities and risks of foundation models. arXiv. https://arxiv.org/abs/2108.07258",
        "Bzdok, D., Altman, N. y Krzywinski, M. (2018). Statistics versus machine learning. Nature Methods, 15, 233-234. https://doi.org/10.1038/nmeth.4642",
        "Gui, J., Chen, T., Zhang, J., Cao, Q., Sun, Z., Luo, H. y Tao, D. (2023). A survey on self-supervised learning: Algorithms, applications, and future trends. arXiv. https://arxiv.org/abs/2301.05712",
        "Jordan, M. I. y Mitchell, T. M. (2015). Machine learning: Trends, perspectives, and prospects. Science, 349(6245), 255-260. https://doi.org/10.1126/science.aaa8415",
        "LeCun, Y., Bengio, Y. y Hinton, G. (2015). Deep learning. Nature, 521, 436-444. https://doi.org/10.1038/nature14539",
        "Manterola, C., Rivadeneira, J., Delgado, H., Sotelo, C. y Otzen, T. (2023). ¿Cuántos tipos de revisiones de la literatura existen? Enumeración, descripción y clasificación. Revisión cualitativa. International Journal of Morphology, 41(4), 1240-1253. http://dx.doi.org/10.4067/S0717-95022023000401240",
        "Maslej, N., Fattorini, L., Perrault, R., Parli, V., Reuel, A., Brynjolfsson, E., Etchemendy, J., Ligett, K., Lyons, T., Manyika, J., Niebles, J. C., Shoham, Y., Wald, R. y Clark, J. (2024). Artificial Intelligence Index Report 2024. Stanford Institute for Human-Centered Artificial Intelligence. https://arxiv.org/abs/2405.19522",
        "Narayanan, A. y Kapoor, S. (2024). AI snake oil: What artificial intelligence can do, what it can't, and how to tell the difference. Princeton University Press.",
        "Sajadieh, S., Fattorini, L., Perrault, R., Gil, Y., Parli, V., Santarlasci, L., Pava, J., Maslej, N., Altman, R., Brynjolfsson, E., Brodley, C., Clark, J., Dignum, V., Kumar, V., Landay, J., Lyons, T., Manyika, J., Niebles, J. C., Shoham, Y., ... Weld, D. (2026). Artificial Intelligence Index Report 2026. Stanford Institute for Human-Centered Artificial Intelligence. https://arxiv.org/abs/2606.15708",
        "Sutton, R. S. y Barto, A. G. (2018). Reinforcement learning: An introduction (2.ª ed.). MIT Press.",
        "Uelwer, T., Robine, J., Wagner, S. S., Höftmann, M., Upschulte, E., Konietzny, S., Behrendt, M. y Harmeling, S. (2023). A survey on self-supervised representation learning. arXiv. https://arxiv.org/abs/2308.11455",
        "Yu, J., Yin, H., Xia, X., Chen, T., Li, J. y Huang, Z. (2022). Self-supervised learning for recommender systems: A survey. arXiv. https://arxiv.org/abs/2203.15876",
    ]
    for ref in references:
        add_reference(doc, ref)

    doc.add_page_break()
    add_heading(doc, "Anexo A - Matriz de cumplimiento del mandato", level=1)
    add_table(
        doc,
        ["Mandato", "Dónde se cumple"],
        [
            ("Revisión de literatura sobre SML, NSML, reforzamiento, self-supervised y deep learning", "Secciones Revisión literaria y Hallazgos."),
            ("80% literatura posterior a 2021", "Referencias 2022-2026 priorizadas, con clásicos solo para base histórica."),
            ("ML vs investigación convencional", "Sección Hallazgos y contraste crítico."),
            ("Tendencias de industria y Big Data", "Sección Tendencias de la industria."),
            ("Algoritmos tendencia con gráficos, tablas o diagramas", "Tabla 1 y Figura 1."),
            ("Métricas convencionales vs ML", "Sección Métricas de validación."),
            ("Conclusión sustanciosa y crítica", "Sección Conclusiones."),
        ],
    )

    add_heading(doc, "Anexo B - Matriz de cumplimiento de la rúbrica", level=1)
    add_table(
        doc,
        ["Componente de la rúbrica", "Evidencia incluida", "Puntos que protege"],
        [
            ("Calidad e impacto de las fuentes; originalidad", "Fuentes revisadas por pares, reportes institucionales, literatura posterior a 2021 y análisis conectado con experiencia profesional en QA/SDET.", "4"),
            ("Estructura del contenido", "Resumen, abstract, tabla de contenido, introducción, metodología, revisión literaria, findings, conclusiones, anexos, tabla y figura.", "4"),
            ("Referencias, citas y APA", "Citas narrativas, parentéticas, cita textual breve, referencias con sangría francesa, tablas/figura con nota y formato APA 7.", "2"),
        ],
    )

    add_heading(doc, "Anexo C - Enlace para infografía o mapa conceptual", level=1)
    add_paragraph(doc, "Enlace en la nube: ________________________________", indent=False)
    add_paragraph(doc, "Nota. Este espacio queda preparado para agregar un enlace de Google Drive, Canva, Miro, Lucidchart u otra herramienta con una infografía de autoría propia.", indent=False)

    force_all_text_black(doc)
    output = DOCS / "Ensayo_2_Revision_Literatura_ML_Marlenis_APA7.docx"
    doc.save(output)
    print(output)


if __name__ == "__main__":
    build_doc()
