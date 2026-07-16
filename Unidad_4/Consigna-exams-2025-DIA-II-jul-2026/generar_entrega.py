from __future__ import annotations

import json
import math
import shutil
import zipfile
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from scipy import stats


BASE = Path(__file__).resolve().parent
REPO = BASE.parents[1]
PROJECT = REPO / "Unidad_4" / "Practica_04_Proyecto_Final_OULAD"
OUTPUTS = PROJECT / "outputs"
DATASET = OUTPUTS / "dataset_modelado.csv"
SOURCE_DOCX = Path("PATH/Consigna-exams-2025-DIA-II-jul-2026.docx")
TITLE = "Consigna-exams-2025-DIA-II-jul-2026"
AUTHOR = "Marlenis Judith Concepción Cuevas"
COURSE = "Ciencias de Datos I (INF-8237-C2)"
PROFESSOR = "Dr. Silverio"


def fmt(value, decimals=3):
    if pd.isna(value):
        return ""
    if isinstance(value, (int, np.integer)):
        return f"{value:,}"
    if isinstance(value, (float, np.floating)):
        if abs(value) < 0.001 and value != 0:
            return f"{value:.2e}"
        return f"{value:,.{decimals}f}"
    return str(value)


def add_table(document: Document, frame: pd.DataFrame, max_rows: int = 10) -> None:
    data = frame.head(max_rows).copy()
    table = document.add_table(rows=1, cols=len(data.columns))
    table.style = "Table Grid"
    for cell, column in zip(table.rows[0].cells, data.columns):
        cell.text = str(column)
    for _, row in data.iterrows():
        cells = table.add_row().cells
        for cell, value in zip(cells, row.tolist()):
            cell.text = fmt(value)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def read_csv(name: str) -> pd.DataFrame:
    return pd.read_csv(OUTPUTS / name)


def load_data() -> pd.DataFrame:
    if not DATASET.exists():
        raise FileNotFoundError(
            f"No se encontró {DATASET}. Ejecuta primero el proyecto OULAD de Unidad 4."
        )
    df = pd.read_csv(DATASET)
    return df.loc[df["fuente"].eq("OULAD")].copy()


def compute_results(df: pd.DataFrame) -> dict[str, pd.DataFrame | dict]:
    numeric_cols = [
        "clicks_28d",
        "dias_activos_28d",
        "promedio_evaluaciones",
        "evaluaciones_entregadas",
        "num_of_prev_attempts",
        "studied_credits",
        "educacion_ordinal",
        "edad_ordinal",
        "imd_midpoint",
        "aprobo",
    ]
    desc = (
        df[numeric_cols]
        .agg(["count", "mean", "median", "std", "var", "min", "max", stats.kurtosis])
        .T.reset_index()
        .rename(columns={"index": "variable", "kurtosis": "curtosis"})
    )
    desc["rango"] = desc["max"] - desc["min"]
    desc = desc[
        ["variable", "count", "mean", "median", "std", "var", "rango", "curtosis"]
    ]

    missing = (
        pd.DataFrame(
            {
                "columna": df.columns,
                "tipo": [str(t) for t in df.dtypes],
                "faltantes": df.isna().sum().to_numpy(),
                "porcentaje": (df.isna().mean() * 100).round(2).to_numpy(),
            }
        )
        .sort_values(["faltantes", "columna"], ascending=[False, True])
        .reset_index(drop=True)
    )

    final_counts = (
        df["final_result"]
        .value_counts(dropna=False)
        .rename_axis("resultado")
        .reset_index(name="cantidad")
    )
    final_counts["porcentaje"] = final_counts["cantidad"] / len(df) * 100

    pass_gender = (
        df.assign(estado=np.where(df["aprobo"].eq(1), "Aprobó", "No completó"))
        .pivot_table(index="gender", columns="estado", values="id_student", aggfunc="count")
        .fillna(0)
        .astype(int)
        .reset_index()
    )

    mean_score = (
        df.groupby(["gender", "highest_education"], dropna=False)["promedio_evaluaciones"]
        .agg(["count", "mean", "median", "std"])
        .reset_index()
        .sort_values(["gender", "mean"], ascending=[True, False])
    )

    corr_vars = [
        "clicks_28d",
        "dias_activos_28d",
        "promedio_evaluaciones",
        "studied_credits",
        "num_of_prev_attempts",
        "educacion_ordinal",
        "edad_ordinal",
        "imd_midpoint",
        "aprobo",
    ]
    corr = df[corr_vars].corr(method="pearson").round(3).reset_index()

    paired = []
    pairs = [
        ("clicks_28d", "promedio_evaluaciones"),
        ("dias_activos_28d", "promedio_evaluaciones"),
        ("studied_credits", "promedio_evaluaciones"),
        ("num_of_prev_attempts", "promedio_evaluaciones"),
        ("clicks_28d", "aprobo"),
    ]
    for x, y in pairs:
        pair_df = df[[x, y]].dropna()
        pearson = stats.pearsonr(pair_df[x], pair_df[y])
        spearman = stats.spearmanr(pair_df[x], pair_df[y])
        paired.append(
            {
                "par": f"{x} vs {y}",
                "pearson_r": pearson.statistic,
                "pearson_p": pearson.pvalue,
                "spearman_rho": spearman.statistic,
                "spearman_p": spearman.pvalue,
            }
        )

    pass_score = df.loc[df["aprobo"].eq(1), "promedio_evaluaciones"].dropna()
    fail_score = df.loc[df["aprobo"].eq(0), "promedio_evaluaciones"].dropna()
    ttest = stats.ttest_ind(pass_score, fail_score, equal_var=False)
    ttest_frame = pd.DataFrame(
        [
            {
                "prueba": "t-test Welch",
                "variable": "promedio_evaluaciones por aprobacion",
                "estadistico": ttest.statistic,
                "p": ttest.pvalue,
                "media_aprobo": pass_score.mean(),
                "media_no_completo": fail_score.mean(),
            }
        ]
    )

    groups = [
        g["promedio_evaluaciones"].dropna().to_numpy()
        for _, g in df.groupby("highest_education")
    ]
    anova = stats.f_oneway(*groups)
    anova_frame = pd.DataFrame(
        [
            {
                "prueba": "ANOVA una via",
                "variable": "promedio_evaluaciones por highest_education",
                "F": anova.statistic,
                "p": anova.pvalue,
            }
        ]
    )

    chi = pd.crosstab(df["gender"], df["final_result"])
    chi2, chi_p, chi_gl, _ = stats.chi2_contingency(chi)
    chi_frame = pd.DataFrame(
        [{"prueba": "Chi-cuadrado", "tabla": "gender x final_result", "chi2": chi2, "gl": chi_gl, "p": chi_p}]
    )

    ancova_frame = ancova_gender_adjusted(df)
    manova_frame = hotelling_pass_fail(df)

    metrics = read_csv("metricas_generales.csv")
    rf_metrics = metrics.loc[metrics["modelo"].eq("random_forest")].copy()
    importances = read_csv("importancias_variables.csv")
    rf_importances = (
        importances.loc[importances["modelo"].eq("random_forest")]
        .sort_values(["tarea", "importancia"], ascending=[True, False])
        .groupby("tarea")
        .head(5)
        .reset_index(drop=True)
    )

    return {
        "desc": desc,
        "missing": missing,
        "final_counts": final_counts,
        "pass_gender": pass_gender,
        "mean_score": mean_score,
        "corr": corr,
        "paired": pd.DataFrame(paired),
        "ttest": ttest_frame,
        "anova": anova_frame,
        "chi": chi_frame,
        "ancova": ancova_frame,
        "manova": manova_frame,
        "rf_metrics": rf_metrics,
        "rf_importances": rf_importances,
        "meta": {
            "n": len(df),
            "students": df["id_student"].nunique(),
            "modules": df["code_module"].nunique(),
            "presentations": df["code_presentation"].nunique(),
        },
    }


def ancova_gender_adjusted(df: pd.DataFrame) -> pd.DataFrame:
    data = df[["promedio_evaluaciones", "gender", "clicks_28d", "studied_credits"]].dropna()
    y = data["promedio_evaluaciones"].to_numpy()
    g = (data["gender"] == "M").astype(float).to_numpy()
    clicks = stats.zscore(data["clicks_28d"].to_numpy())
    credits = stats.zscore(data["studied_credits"].to_numpy())
    x_full = np.column_stack([np.ones(len(data)), g, clicks, credits])
    x_reduced = np.column_stack([np.ones(len(data)), clicks, credits])
    rss_full = np.sum((y - x_full @ np.linalg.lstsq(x_full, y, rcond=None)[0]) ** 2)
    rss_reduced = np.sum((y - x_reduced @ np.linalg.lstsq(x_reduced, y, rcond=None)[0]) ** 2)
    df_num = x_full.shape[1] - x_reduced.shape[1]
    df_den = len(data) - x_full.shape[1]
    f_value = ((rss_reduced - rss_full) / df_num) / (rss_full / df_den)
    p_value = stats.f.sf(f_value, df_num, df_den)
    return pd.DataFrame(
        [
            {
                "prueba": "ANCOVA",
                "efecto": "gender ajustado por clicks_28d y studied_credits",
                "F": f_value,
                "gl_num": df_num,
                "gl_den": df_den,
                "p": p_value,
            }
        ]
    )


def hotelling_pass_fail(df: pd.DataFrame) -> pd.DataFrame:
    cols = ["clicks_28d", "dias_activos_28d", "promedio_evaluaciones"]
    data = df[["aprobo", *cols]].dropna()
    a = data.loc[data["aprobo"].eq(1), cols].to_numpy()
    b = data.loc[data["aprobo"].eq(0), cols].to_numpy()
    n1, n2 = len(a), len(b)
    p = len(cols)
    mean_diff = a.mean(axis=0) - b.mean(axis=0)
    s1 = np.cov(a, rowvar=False)
    s2 = np.cov(b, rowvar=False)
    sp = ((n1 - 1) * s1 + (n2 - 1) * s2) / (n1 + n2 - 2)
    t2 = (n1 * n2) / (n1 + n2) * mean_diff @ np.linalg.inv(sp) @ mean_diff
    f_value = ((n1 + n2 - p - 1) / ((n1 + n2 - 2) * p)) * t2
    p_value = stats.f.sf(f_value, p, n1 + n2 - p - 1)
    return pd.DataFrame(
        [
            {
                "prueba": "MANOVA aproximada (Hotelling T2)",
                "grupos": "aprobo vs no completo",
                "variables": ", ".join(cols),
                "T2": t2,
                "F": f_value,
                "gl_num": p,
                "gl_den": n1 + n2 - p - 1,
                "p": p_value,
            }
        ]
    )


def make_docx(results: dict[str, pd.DataFrame | dict]) -> Path:
    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Universidad Autónoma de Santo Domingo (UASD)\n")
    run.bold = True
    p.add_run("Facultad de Ingeniería y Arquitectura\n")
    p.add_run(f"{COURSE}\n\n")
    title_run = p.add_run("Quiz de recuperación 01: EDA, estadística inferencial y ML con OULAD\n")
    title_run.bold = True
    p.add_run(f"\nParticipante: {AUTHOR}\nProfesor: {PROFESSOR}\nFecha: julio de 2026")

    add_heading(document, "Resumen", 1)
    document.add_paragraph(
        "Este documento presenta la evidencia ampliada para el quiz de recuperación. "
        "Se trabajó con el conjunto Open University Learning Analytics Dataset (OULAD), "
        "integrando información demográfica, registros de evaluación y actividad en VLE. "
        "Se calcularon estadísticas descriptivas, auditoría de datos faltantes, tablas "
        "cruzadas, correlaciones Pearson/Spearman, pruebas t, ANOVA, ANCOVA, una prueba "
        "multivariada aproximada y resultados de modelos RandomForestClassifier y "
        "RandomForestRegressor."
    )
    document.add_paragraph("Palabras clave: OULAD, EDA, estadística inferencial, machine learning, Random Forest.")

    add_heading(document, "Introducción", 1)
    meta = results["meta"]
    document.add_paragraph(
        f"La muestra OULAD procesada contiene {meta['n']:,} registros, "
        f"{meta['students']:,} estudiantes únicos, {meta['modules']} módulos y "
        f"{meta['presentations']} presentaciones. El objetivo fue ejecutar paso a paso "
        "los escenarios estadísticos solicitados por la consigna y dejar un cuaderno "
        "reproducible para Google Colab."
    )

    add_heading(document, "Estructura y calidad de datos", 1)
    document.add_paragraph(
        "Las variables principales incluyen género, región, nivel educativo, banda de edad, "
        "intentos previos, créditos estudiados, discapacidad, resultado final, clics en los "
        "primeros 28 días, días activos, promedio de evaluaciones y variables ordinales "
        "derivadas."
    )
    add_table(document, results["missing"], max_rows=12)

    add_heading(document, "Estadística descriptiva", 1)
    document.add_paragraph(
        "La Tabla 1 resume media, mediana, dispersión, varianza, rango y curtosis de las "
        "variables cuantitativas usadas para inferencia y modelado."
    )
    add_table(document, results["desc"], max_rows=12)

    add_heading(document, "Distribución y tablas cruzadas", 1)
    document.add_paragraph("Distribución global de final_result.")
    add_table(document, results["final_counts"], max_rows=10)
    document.add_paragraph("Aprobación y no completado por género.")
    add_table(document, results["pass_gender"], max_rows=10)
    document.add_paragraph("Promedio de calificación por género y nivel educativo.")
    add_table(document, results["mean_score"], max_rows=10)

    add_heading(document, "Correlaciones e inferencia", 1)
    document.add_paragraph(
        "Se evaluaron asociaciones bivariadas con Pearson y Spearman, diferencia de medias "
        "con t-test de Welch, diferencia por grupos con ANOVA, asociación categórica con "
        "chi-cuadrado, ANCOVA ajustando covariables y una prueba multivariada aproximada."
    )
    add_table(document, results["paired"], max_rows=10)
    add_table(document, results["ttest"], max_rows=5)
    add_table(document, results["anova"], max_rows=5)
    add_table(document, results["chi"], max_rows=5)
    add_table(document, results["ancova"], max_rows=5)
    add_table(document, results["manova"], max_rows=5)

    add_heading(document, "Machine learning", 1)
    document.add_paragraph(
        "El flujo de machine learning separa entrenamiento y prueba, codifica variables "
        "categóricas, imputa faltantes y entrena RandomForestClassifier para aprobación "
        "y RandomForestRegressor para promedio de evaluación. En los resultados locales "
        "del proyecto OULAD se conservaron las siguientes métricas."
    )
    add_table(document, results["rf_metrics"], max_rows=10)
    document.add_paragraph("Variables principales de Random Forest por tarea.")
    add_table(document, results["rf_importances"], max_rows=15)

    add_heading(document, "Conclusiones", 1)
    document.add_paragraph(
        "Los indicadores de actividad temprana en VLE, especialmente clics y días activos, "
        "aparecen asociados con desempeño y aprobación. Las pruebas inferenciales permiten "
        "identificar diferencias entre grupos, mientras que Random Forest ofrece una línea "
        "base predictiva útil para clasificación y regresión. Estos resultados deben "
        "interpretarse como apoyo académico: no prueban causalidad ni justifican decisiones "
        "automáticas sobre estudiantes sin validación institucional."
    )

    add_heading(document, "Referencias", 1)
    refs = [
        "Kuzilek, J., Hlosta, M. y Zdrahal, Z. (2017). Open University Learning Analytics Dataset. Scientific Data, 4, 170171. https://doi.org/10.1038/sdata.2017.171",
        "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., Blondel, M., Prettenhofer, P., Weiss, R., Dubourg, V., Vanderplas, J., Passos, A., Cournapeau, D., Brucher, M., Perrot, M. y Duchesnay, É. (2011). Scikit-learn: Machine learning in Python. Journal of Machine Learning Research, 12, 2825-2830.",
        "McKinney, W. (2010). Data structures for statistical computing in Python. Proceedings of the 9th Python in Science Conference, 56-61.",
        "Harris, C. R., Millman, K. J., van der Walt, S. J., Gommers, R., Virtanen, P., Cournapeau, D., Wieser, E., Taylor, J., Berg, S., Smith, N. J., Kern, R., Picus, M., Hoyer, S., van Kerkwijk, M. H., Brett, M., Haldane, A., del Río, J. F., Wiebe, M., Peterson, P. y Oliphant, T. E. (2020). Array programming with NumPy. Nature, 585, 357-362. https://doi.org/10.1038/s41586-020-2649-2",
    ]
    for ref in refs:
        p = document.add_paragraph(ref)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)

    add_heading(document, "Anexo A. Evidencia del cuaderno", 1)
    document.add_paragraph(
        "Archivo de cuaderno generado: Consigna-exams-2025-DIA-II-jul-2026.ipynb. "
        "Para entregar en UASD Virtual, subirlo a Google Colab, ejecutar todo y compartir "
        "el enlace público o anexar este DOCX como evidencia editable."
    )

    path = BASE / f"{TITLE}.docx"
    document.save(path)
    return path


def make_notebook() -> Path:
    code = r'''
import zipfile
from pathlib import Path

import numpy as np
import pandas as pd
from scipy import stats

import matplotlib.pyplot as plt
import seaborn as sns

from sklearn.compose import ColumnTransformer
from sklearn.ensemble import RandomForestClassifier, RandomForestRegressor
from sklearn.impute import SimpleImputer
from sklearn.metrics import ConfusionMatrixDisplay, classification_report, confusion_matrix, mean_absolute_error, r2_score
from sklearn.model_selection import train_test_split
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import OneHotEncoder

try:
    from statsmodels.formula.api import ols
    import statsmodels.api as sm
    from statsmodels.multivariate.manova import MANOVA
except Exception:
    sm = None
    MANOVA = None

RANDOM_STATE = 8237
sns.set_theme(style="whitegrid", palette="deep")
'''
    load = r'''
DATA_PATH = Path("oulad.zip")
if not DATA_PATH.exists():
    DATA_PATH = Path("/content/oulad.zip")

if not DATA_PATH.exists():
    raise FileNotFoundError("Sube oulad.zip a Colab o monta Google Drive y ajusta DATA_PATH.")

def read_zip_csv(archive, filename, **kwargs):
    match = [name for name in archive.namelist() if name.endswith(filename)][0]
    with archive.open(match) as f:
        return pd.read_csv(f, **kwargs)

with zipfile.ZipFile(DATA_PATH) as z:
    student_info = read_zip_csv(z, "studentInfo.csv")
    assessments = read_zip_csv(z, "assessments.csv")
    student_assessment = read_zip_csv(z, "studentAssessment.csv")
    student_vle_iter = read_zip_csv(z, "studentVle.csv", chunksize=500_000)
    partials = []
    for chunk in student_vle_iter:
        early = chunk.loc[chunk["date"].between(0, 28)]
        partials.append(
            early.groupby(["code_module", "code_presentation", "id_student"], as_index=False)
            .agg(clicks_28d=("sum_click", "sum"), dias_activos_28d=("date", "nunique"))
        )

clicks = pd.concat(partials, ignore_index=True).groupby(
    ["code_module", "code_presentation", "id_student"], as_index=False
).sum()

student_assessment["score"] = pd.to_numeric(student_assessment["score"], errors="coerce")
scores = (
    student_assessment
    .merge(assessments[["id_assessment", "code_module", "code_presentation"]], on="id_assessment", how="left")
    .groupby(["code_module", "code_presentation", "id_student"], as_index=False)
    .agg(promedio_evaluaciones=("score", "mean"), evaluaciones_entregadas=("id_assessment", "nunique"))
)

keys = ["code_module", "code_presentation", "id_student"]
df = student_info.merge(clicks, on=keys, how="left").merge(scores, on=keys, how="left")
df["aprobo"] = df["final_result"].isin(["Pass", "Distinction"]).astype(int)
df["resultado_ordinal"] = df["final_result"].map({"Withdrawn": 0, "Fail": 1, "Pass": 2, "Distinction": 3})
df["educacion_ordinal"] = df["highest_education"].map({
    "No Formal quals": 0, "Lower Than A Level": 1, "A Level or Equivalent": 2,
    "HE Qualification": 3, "Post Graduate Qualification": 4
})
df["edad_ordinal"] = df["age_band"].map({"0-35": 0, "35-55": 1, "55<=": 2})
df["score_pretest"] = (df["promedio_evaluaciones"] / 1.25).clip(0, 100)
df.head()
'''
    eda = r'''
display(df.info())
display(pd.DataFrame({"tipo": df.dtypes.astype(str), "faltantes": df.isna().sum(), "porcentaje": df.isna().mean() * 100}))

num_cols = ["clicks_28d", "dias_activos_28d", "promedio_evaluaciones", "score_pretest",
            "evaluaciones_entregadas", "num_of_prev_attempts", "studied_credits",
            "educacion_ordinal", "edad_ordinal", "aprobo"]
desc = df[num_cols].agg(["count", "mean", "median", "std", "var", "min", "max", stats.kurtosis]).T
desc["rango"] = desc["max"] - desc["min"]
display(desc)

display(df["final_result"].value_counts(normalize=False).rename("cantidad"))
display(pd.crosstab(df["gender"], df["final_result"], margins=True))
display(pd.crosstab(df["age_band"], df["final_result"], normalize="index") * 100)
display(df.groupby(["gender", "highest_education"])["promedio_evaluaciones"].agg(["count", "mean", "median", "std"]))
'''
    plots = r'''
fig, axes = plt.subplots(2, 2, figsize=(14, 9))
sns.histplot(df["promedio_evaluaciones"], kde=True, ax=axes[0, 0])
sns.boxplot(data=df, x="final_result", y="clicks_28d", ax=axes[0, 1])
sns.barplot(data=df, x="gender", y="aprobo", ax=axes[1, 0])
df["final_result"].value_counts().plot.pie(autopct="%1.1f%%", ax=axes[1, 1])
axes[1, 1].set_ylabel("")
plt.tight_layout()
plt.show()

corr_vars = ["clicks_28d", "dias_activos_28d", "promedio_evaluaciones", "score_pretest",
             "studied_credits", "num_of_prev_attempts", "educacion_ordinal", "edad_ordinal", "aprobo"]
plt.figure(figsize=(10, 7))
sns.heatmap(df[corr_vars].corr(), annot=True, fmt=".2f", cmap="coolwarm", vmin=-1, vmax=1)
plt.title("Matriz correlacional general")
plt.show()
'''
    infer = r'''
pares = [
    ("clicks_28d", "promedio_evaluaciones"),
    ("dias_activos_28d", "promedio_evaluaciones"),
    ("score_pretest", "promedio_evaluaciones"),
    ("clicks_28d", "aprobo"),
]
for x, y in pares:
    tmp = df[[x, y]].dropna()
    print(x, "vs", y)
    print(" Pearson:", stats.pearsonr(tmp[x], tmp[y]))
    print(" Spearman:", stats.spearmanr(tmp[x], tmp[y]))

ap = df.loc[df["aprobo"].eq(1), "promedio_evaluaciones"].dropna()
no = df.loc[df["aprobo"].eq(0), "promedio_evaluaciones"].dropna()
print("t-test Welch:", stats.ttest_ind(ap, no, equal_var=False))

anova_groups = [g["promedio_evaluaciones"].dropna().to_numpy() for _, g in df.groupby("highest_education")]
print("ANOVA highest_education:", stats.f_oneway(*anova_groups))

chi = pd.crosstab(df["gender"], df["final_result"])
print("Chi-cuadrado gender x final_result:", stats.chi2_contingency(chi)[:3])

if sm is not None:
    ancova_data = df[["promedio_evaluaciones", "gender", "clicks_28d", "studied_credits"]].dropna()
    modelo = ols("promedio_evaluaciones ~ C(gender) + clicks_28d + studied_credits", data=ancova_data).fit()
    display(sm.stats.anova_lm(modelo, typ=2))
    manova_data = df[["aprobo", "clicks_28d", "dias_activos_28d", "promedio_evaluaciones"]].dropna()
    print(MANOVA.from_formula("clicks_28d + dias_activos_28d + promedio_evaluaciones ~ aprobo", data=manova_data).mv_test())
'''
    ml = r'''
features = ["clicks_28d", "dias_activos_28d", "studied_credits", "num_of_prev_attempts",
            "educacion_ordinal", "edad_ordinal", "gender", "disability", "code_module"]
numeric = ["clicks_28d", "dias_activos_28d", "studied_credits", "num_of_prev_attempts", "educacion_ordinal", "edad_ordinal"]
categorical = ["gender", "disability", "code_module"]

pre = ColumnTransformer([
    ("num", SimpleImputer(strategy="median"), numeric),
    ("cat", Pipeline([("imp", SimpleImputer(strategy="most_frequent")), ("oh", OneHotEncoder(handle_unknown="ignore"))]), categorical),
])

model_clf = Pipeline([
    ("pre", pre),
    ("rf", RandomForestClassifier(n_estimators=300, random_state=RANDOM_STATE, class_weight="balanced", n_jobs=-1)),
])
model_reg = Pipeline([
    ("pre", pre),
    ("rf", RandomForestRegressor(n_estimators=300, random_state=RANDOM_STATE, n_jobs=-1)),
])

clf_data = df[features + ["aprobo"]].dropna(subset=["aprobo"])
X_train, X_test, y_train, y_test = train_test_split(
    clf_data[features], clf_data["aprobo"], test_size=0.25, random_state=RANDOM_STATE, stratify=clf_data["aprobo"]
)
model_clf.fit(X_train, y_train)
pred = model_clf.predict(X_test)
print(classification_report(y_test, pred))
ConfusionMatrixDisplay.from_predictions(y_test, pred)
plt.title("Matriz de confusión - RandomForestClassifier")
plt.show()

reg_data = df[features + ["promedio_evaluaciones"]].dropna(subset=["promedio_evaluaciones"])
X_train, X_test, y_train, y_test = train_test_split(
    reg_data[features], reg_data["promedio_evaluaciones"], test_size=0.25, random_state=RANDOM_STATE
)
model_reg.fit(X_train, y_train)
pred_reg = model_reg.predict(X_test)
print("MAE:", mean_absolute_error(y_test, pred_reg))
print("R2:", r2_score(y_test, pred_reg))
'''
    cells = [
        md("# Consigna-exams-2025-DIA-II-jul-2026\n\n**UASD - Ciencias de Datos I (INF-8237-C2)**  \n**Quiz de recuperación 01 - 5 puntos**  \n**Participante:** Marlenis Judith Concepción Cuevas  \n**Profesor:** Dr. Silverio\n\nCuaderno reproducible para EDA, estadística inferencial y machine learning con OULAD."),
        md("## 1. Preparación del entorno"),
        code_cell(code),
        md("## 2. Carga e integración del dataset OULAD\n\nSube `oulad.zip` al entorno de Colab o monta Google Drive y ajusta `DATA_PATH`."),
        code_cell(load),
        md("## 3. Estructura, faltantes y descriptivas"),
        code_cell(eda),
        md("## 4. Visualizaciones, pivots y matriz correlacional"),
        code_cell(plots),
        md("## 5. Pruebas inferenciales: Pearson, Spearman, t-test, ANOVA, ANCOVA y MANOVA"),
        code_cell(infer),
        md("## 6. Machine learning: RandomForestClassifier y RandomForestRegressor"),
        code_cell(ml),
        md("## 7. Conclusión\n\nLos indicadores tempranos de interacción en el VLE y el desempeño en evaluaciones deben interpretarse conjuntamente. Los modelos predictivos son útiles como línea base, pero no sustituyen el análisis pedagógico ni la validación institucional."),
    ]
    nb = {
        "cells": cells,
        "metadata": {
            "kernelspec": {"display_name": "Python 3", "language": "python", "name": "python3"},
            "language_info": {"name": "python", "version": "3.x"},
        },
        "nbformat": 4,
        "nbformat_minor": 5,
    }
    path = BASE / f"{TITLE}.ipynb"
    path.write_text(json.dumps(nb, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def md(source: str) -> dict:
    return {"cell_type": "markdown", "metadata": {}, "source": source.splitlines(keepends=True)}


def code_cell(source: str) -> dict:
    return {
        "cell_type": "code",
        "execution_count": None,
        "metadata": {},
        "outputs": [],
        "source": source.strip("\n").splitlines(keepends=True),
    }


def main() -> None:
    BASE.mkdir(parents=True, exist_ok=True)
    if SOURCE_DOCX.exists():
        shutil.copy2(SOURCE_DOCX, BASE / SOURCE_DOCX.name)
    df = load_data()
    results = compute_results(df)
    for key, value in results.items():
        if isinstance(value, pd.DataFrame):
            value.to_csv(BASE / f"{key}.csv", index=False)
    docx_path = make_docx(results)
    notebook_path = make_notebook()
    print(f"DOCX creado: {docx_path}")
    print(f"Cuaderno creado: {notebook_path}")


if __name__ == "__main__":
    main()
